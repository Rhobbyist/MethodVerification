from ast import Pass
import xlrd
from report.models import *
from docx import Document
from report.effectnum import *
import re


def PTfileread(files, Detectionplatform, project, platform, manufacturers, digits, ZP_Method_precursor_ion, ZP_Method_product_ion, normAB, Number_of_compounds):

    # 第一步 后台数据抓取（待测物质,可接受标准，单位）
    zqd = Special.objects.get(project=project)
    pt_special = PTspecial.objects.get(special=zqd)
    pt_accept = PTspecialaccept.objects.filter(pTspecial=pt_special)
    PTnorm = []  # 待测物质列表
    PTloq = []  # loq列表,进行数据显示判断
    PTrange1 = []  # 可接受标准一适用范围，与待测物质列表一一对应
    PTstandard1 = []  # 可接受标准一
    PTrange2 = []  # 可接受标准二适用范围，与待测物质列表一一对应
    PTstandard2 = []  # 可接受标准二
    PTunit = Special.objects.get(project=project).unit  # 单位

    for i in pt_accept:
        PTnorm.append(i.norm)
        PTloq.append(i.loq)
        PTrange1.append(i.range1)
        PTstandard1.append(i.accept1)
        PTrange2.append(i.range2)
        PTstandard2.append(i.accept2)

    # 如果没在后台管理系统中设置化合物名称直接返回并提示
    if all(i is None for i in PTnorm):
        error_message = "尚未在后台管理系统中设置PT的化合物名称,请设置后重新提交数据!"
        return {"error_message": error_message}

    # Python判断列表中是否为空，包括None
    # if all(i is None for i in PTrange1):
    #     error_message="尚未在后台管理系统中设置PT的可接受标准,请设置后重新提交数据!"
    #     return {"error_message":error_message}

    # AB厂家,未在后台管理系统里规范设置定量离子对数值,直接返回并提示(含有“定量”关键词的输入框个数与项目中实际化合物个数不一致)
    if manufacturers == "AB":
        if len(normAB) != Number_of_compounds:
            error_message = "未在后台管理系统里规范设置定量离子对数值，请检查并规范设置后重新提交数据!"
            return {"error_message": error_message}

    #  第二步:开始文件读取
    '''
    注释:csv,txt,xlsx,docx 4种格式数据读取完毕后,需要生成一个字典PT_dict,数据格式如下：
    print(PT_dict):
    {'MN': [['PT1', 0.49, '±0.075nmol/L'], ['PT10', 3.32, '±15.0%'], ['PT19', 3.31, '±15.0%'], ['PT28', 3.29, '±15.0%']],
    'NMN': [['PT4', 5.96, '±20.0%'], ['PT13', 4.37, '±20.0%'], ['PT22', 4.38, '±20.0%'], ['PT31', 4.25, '±20.0%']], 
    '3-MT': [['PT7', 8.78, '±30.0%'], ['PT16', 1.38, '±30.0%'], ['PT25', 1.38, '±30.0%'], ['PT34', 1.37, '±30.0%']]}
    '''

    # 头部定义相关需要提取生成的结果
    PT_dict = {}
    for i in range(len(PTnorm)):
        PT_dict[PTnorm[i]] = []

    # 各仪器平台及各仪器厂家数据读取
    for file in files:
        if platform == "液质":
            if manufacturers == "Agilent":
                # 1 读取csv文件
                file.seek(0)  # https://www.jianshu.com/p/0d15ed85df2b
                file_data = file.read().decode('utf-8')
                lines = file_data.split('\r\n')
                for i in range(len(lines)):
                    if len(lines[i]) != 0:
                        # 以逗号分隔字符串,但忽略双引号内的逗号
                        lines[i] = re.split(r',\s*(?![^"]*\"\,)', lines[i])
                        # lines[i]=lines[i].split(',') # 按逗号分隔后把每一行都变成一个列表
                    else:
                        lines[i] = re.split(r',\s*(?![^"]*\"\,)', lines[i])
                        del lines[i]  # 最后一行如为空行，则删除该元素

                # 药物浓度平台的数据导出格式较不同
                if Detectionplatform=="治疗药物检测平台":
                    norm = []  # 化合物列表
                    for j in range(len(lines[0])):  # 从第一行开始
                        if " Results" in lines[0][j]:
                            # 若原始字符串中含有','，切割完后首位会多出一个'"',需去除
                            if lines[0][j].split(" Results")[0][0] != '"':
                                norm.append(lines[0][j].split(" Results")[0])
                            else:
                                norm.append(lines[0][j].split(" Results")[0][1:])

                    # 定性指标的索引
                    dxindex = []
                    for i in range(len(norm)):
                        if "2" in norm[i]:
                            dxindex.append(i)

                    # 删除norm中的定性指标和定量指标的"1"
                    norm=[norm[i] for i in range(len(norm)) if("2" not in norm[i])]
                    norm=[norm[i].split("1")[0] for i in range(len(norm)) if("1" in norm[i])]
                    
                    # 从第二行确定实验号（Sample Name）,浓度（Exp. Conc.）的索引
                    nameindex = 0  # 实验号索引
                    concindex = []  # 浓度索引列表（可能不止一个化合物，因此需要把索引放在一个列表里）
                    for j in range(len(lines[1])):  # 从第二行开始
                        if lines[1][j] == "Name":
                            nameindex = j
                        elif lines[1][j] == "Final Conc.":
                            concindex.append(j)
                    
                    # 删除concindex中的dxindex
                    concindex=[concindex[i] for i in range(len(concindex)) if(i not in dxindex)]
                    print(concindex)
                
                else:
                    # 从第一行确定化合物名称(含有"-Q Results"),并添加进入化合物列表
                    norm = []  # 化合物列表
                    for j in range(len(lines[0])):  # 从第一行开始
                        if "-Q Results" in lines[0][j]:
                            norm.append(lines[0][j].split("-Q")[0])

                    # 从第二行确定实验号（Sample Name）,浓度（Exp. Conc.）的索引
                    nameindex = 0  # 实验号索引
                    concindex = []  # 浓度索引列表（可能不止一个化合物，因此需要把索引放在一个列表里）
                    for j in range(len(lines[1])):  # 从第二行开始
                        if lines[1][j] == "Name":
                            nameindex = j
                        elif lines[1][j] == "Calc. Conc.":
                            concindex.append(j)

                # 未规范设置表格列名
                if concindex == []:
                    error_message = "未规范设置文件列名，请检查并规范设置后重新提交数据!"
                    return {"error_message": error_message}

                # 判断是否在后台管理系统中设置了可接受标准

                # 第一种情况，未设置，前端调用可接受区间模板
                if all(i is None for i in PTrange1):
                    templates = 1
                    for j in range(len(concindex)):
                        for i in range(len(lines)):  # 循环原始数据中的每一行
                            if "PT" in lines[i][nameindex]:

                                # 第一种情况：未稀释，结果不需要除以稀释倍数
                                if "times" not in lines[i][nameindex]:
                                    PT_dict[PTnorm[j]].append([lines[i][nameindex], effectnum(lines[i][concindex[j]], digits)])
                                

                                # 第二种情况：进行稀释，结果需要除以稀释倍数(实验号命名规则：PT-2times-1）
                                else:
                                    dilute = int(lines[i][nameindex][3:4])
                                    actualconc = float(lines[i][concindex[j]])*dilute
                                    PT_dict[PTnorm[j]].append([lines[i][nameindex], effectnum(actualconc, digits)])
            
                # 第二种情况，已设置，前端调用可接受标准模板
                else:
                    templates = 2
                    # 匹配原始数据中与PT相关(实验号前含有"PT")的行
                    for j in range(len(concindex)):
                        for i in range(len(lines)):  # 循环原始数据中的每一行
                            if "PT" in lines[i][nameindex]:

                                # 第一种情况：未稀释，结果不需要除以稀释倍数
                                if "times" not in lines[i][nameindex]:
                                    # 小于range1,添加第一个可接受标准
                                    if float(lines[i][concindex[j]]) < PTrange1[j]:
                                        PT_dict[PTnorm[j]].append([lines[i][nameindex], effectnum(lines[i][concindex[j]], digits), "±"+" "+str(PTstandard1[j])+" "+PTunit])
                                    # 大于range2,添加第二个可接受标准
                                    elif float(lines[i][concindex[j]]) >= PTrange2[j]:
                                        PT_dict[PTnorm[j]].append([lines[i][nameindex], effectnum(lines[i][concindex[j]], digits), "±"+" "+str(PTstandard2[j])+" "+"%"])

                                # 第二种情况：进行稀释，结果需要除以稀释倍数(实验号命名规则：PT-2times-1）
                                else:
                                    dilute = int(lines[i][nameindex][3:4])
                                    actualconc = float(lines[i][concindex[j]])*dilute

                                    if float(actualconc) < PTrange1[j]:
                                        PT_dict[PTnorm[j]].append([lines[i][nameindex], effectnum(actualconc, digits), "±"+" "+str(PTstandard1[j])+" "+PTunit])
                                    elif float(actualconc) >= PTrange2[j]:
                                        PT_dict[PTnorm[j]].append([lines[i][nameindex], effectnum(actualconc, digits), "±"+" "+str(PTstandard2[j])+" "+"%"])

            elif manufacturers == "Waters":
                # 内标标识
                ISlist = ["D3", "D4", "D5", "D6", "D7", "D8", "dx", "d4", "d8"]

                data = xlrd.open_workbook(filename=None, file_contents=file.read())  # 读取表格
                file_data = data.sheets()[0]
                nrows = file_data.nrows
                ncols = file_data.ncols

                norm = []  # 化合物列表
                Compound_row = []  # 含有“Compound”关键词的所在行(包含内标)
                norm_row = []  # 实际化合物所在行(不包含内标)
                for i in range(nrows):
                    if "Compound" in str(file_data.row_values(i)[0]) and ":" in str(file_data.row_values(i)[0]):
                        Compound_row.append(i)

                    # 判断是否含有内标标识
                    if all(j not in str(file_data.row_values(i)[0]) for j in ISlist):
                        # 如果某一行第一列含有关键词"Compound"，则该行中含有化合物名称，化合物名称在：后
                        if "Compound" in str(file_data.row_values(i)[0]) and ":" in str(file_data.row_values(i)[0]):
                            norm.append(file_data.row_values(i)[0].split(":")[1].strip())  # strip()的作用是去除前后空格
                            norm_row.append(i)

                # # 第一种情况，不含有内标
                # if len(Compound_row) == len(norm_row):
                #     pass

                # # 第二种情况，含有内标
                # else:
                #     nrows = Compound_row[len(norm_row)]

                nameindex = 0
                concindex = 0
                # 第一个化合物表格确定samplename和浓度所在列，norm_row[0]为第一个化合物所在行，+2是该化合物表格位于该化合物所在行的下两行
                for i in range(len(file_data.row_values(norm_row[0]+2))):
                    if file_data.row_values(norm_row[0]+2)[i] == "ID":
                        nameindex = i
                    elif "nmol/L" in file_data.row_values(norm_row[0]+2)[i]:
                        concindex = i

                # 未规范设置表格列名
                if nameindex == 0 or concindex == []:
                    error_message = "未规范设置文件列名，请检查并规范设置后重新提交数据!"
                    return {"error_message": error_message}

                # 判断是否在后台管理系统中设置了可接受标准

                # 第一种情况，未设置，前端调用可接受区间模板
                if all(i is None for i in PTrange1):
                    templates = 1
                    for j in range(len(norm)):
                        if j < len(norm)-1:  # 如果不是最后一个化合物，索引为该化合物所在行到后一个化合物所在行
                            for i in range(norm_row[j], norm_row[j+1]):
                                if "PT" in file_data.row_values(i)[nameindex]:

                                    # 第一种情况：未稀释，结果不需要除以稀释倍数
                                    if "times" not in file_data.row_values(i)[nameindex]:
                                        PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex], effectnum(file_data.row_values(i)[concindex], digits)])

                                    # 第二种情况：进行稀释，结果需要除以稀释倍数(实验号命名规则：PT-2times-1）
                                    else:
                                        dilute = int(file_data.row_values(i)[nameindex][3:4])
                                        actualconc = float(file_data.row_values(i)[concindex])*dilute
                                        PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex], effectnum(actualconc, digits)])

                        else:
                            for i in range(norm_row[j], nrows):
                                if "PT" in file_data.row_values(i)[nameindex]:

                                    # 第一种情况：未稀释，结果不需要除以稀释倍数
                                    if "times" not in file_data.row_values(i)[nameindex]:
                                        PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex], effectnum(file_data.row_values(i)[concindex], digits)])

                                    # 第二种情况：进行稀释，结果需要除以稀释倍数(实验号命名规则：PT-2times-1）
                                    else:
                                        dilute = int(file_data.row_values(i)[nameindex][3:4])
                                        actualconc = float(file_data.row_values(i)[concindex])*dilute
                                        PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex], effectnum(actualconc, digits)])

                # 第二种情况，已设置，前端调用可接受标准模板
                else:
                    templates = 2
                    for j in range(len(norm)):
                        if j < len(norm)-1:  # 如果不是最后一个化合物，索引为该化合物所在行到后一个化合物所在行
                            for i in range(norm_row[j], norm_row[j+1]):
                                if "PT" in file_data.row_values(i)[nameindex]:

                                    # 第一种情况：未稀释，结果不需要除以稀释倍数
                                    if "times" not in file_data.row_values(i)[nameindex]:
                                        # 小于range1,添加第一个可接受标准
                                        if float(file_data.row_values(i)[concindex]) < PTrange1[j]:
                                            PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex], effectnum(file_data.row_values(i)[concindex], digits), "±"+" "+str(PTstandard1[j])+" "+PTunit])
                                        # 大于range2,添加第二个可接受标准
                                        elif float(file_data.row_values(i)[concindex]) >= PTrange2[j]:
                                            PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex], effectnum(file_data.row_values(i)[concindex], digits), "±"+" "+str(PTstandard2[j])+" "+"%"])

                                    # 第二种情况：进行稀释，结果需要除以稀释倍数(实验号命名规则：PT-2times-1）
                                    else:
                                        dilute = int(file_data.row_values(i)[nameindex][3:4])
                                        actualconc = float(file_data.row_values(i)[concindex])*dilute

                                        if float(actualconc) < PTrange1[j]:
                                            PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex], effectnum(actualconc, digits), "±"+" "+str(PTstandard1[j])+" "+PTunit])
                                        elif float(actualconc) >= PTrange2[j]:
                                            PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex], effectnum(actualconc, digits), "±"+" "+str(PTstandard2[j])+" "+"%"])

                        else:  # 如果是最后一个化合物，索引为该化合物所在行到总行数
                            for i in range(norm_row[j], nrows):
                                if "PT" in file_data.row_values(i)[nameindex]:

                                    # 第一种情况：未稀释，结果不需要除以稀释倍数
                                    if "times" not in file_data.row_values(i)[nameindex]:
                                        # 小于range1,添加第一个可接受标准
                                        if float(file_data.row_values(i)[concindex]) < PTrange1[j]:
                                            PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex], effectnum(file_data.row_values(i)[concindex], digits), "±"+" "+str(PTstandard1[j])+" "+PTunit])
                                        # 大于range2,添加第二个可接受标准
                                        elif float(file_data.row_values(i)[concindex]) >= PTrange2[j]:
                                            PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex], effectnum(file_data.row_values(i)[concindex], digits), "±"+" "+str(PTstandard2[j])+" "+"%"])

                                    # 第二种情况：进行稀释，结果需要除以稀释倍数(实验号命名规则：PT-2times-1）
                                    else:
                                        dilute = int(file_data.row_values(i)[nameindex][3:4])
                                        actualconc = float(file_data.row_values(i)[concindex])*dilute

                                        if float(actualconc) < PTrange1[j]:
                                            PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex], effectnum(actualconc, digits), "±"+" "+str(PTstandard1[j])+" "+PTunit])
                                        elif float(actualconc) >= PTrange2[j]:
                                            PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex], effectnum(actualconc, digits), "±"+" "+str(PTstandard2[j])+" "+"%"])

            elif manufacturers == "Thermo":
                data = xlrd.open_workbook(filename=None, file_contents=file.read())  # 读取表格
                norm = []  # Thermo的原始数据格式为一个化合物一个sheet,获取每个sheet的名字,与PTnorm相等的即为需要的sheet
                sheetindex = []  # 需要的化合物所在sheet索引列表
                for index in range(len(data.sheet_names())):
                    norm.append(data.sheet_names()[index].split("Sheet1")[1])
                    sheetindex.append(index)

                # 循环读取每个sheet工作表,即为每个化合物的表
                for j in range(len(sheetindex)):
                    file_data = data.sheets()[sheetindex[j]]
                    nrows = file_data.nrows
                    ncols = file_data.ncols

                    # 第一行确定samplename和浓度所在列
                    nameindex = 0
                    concindex = 0
                    for i in range(len(file_data.row_values(8))):
                        print(file_data.row_values(8)[i])
                        if file_data.row_values(8)[i] == "Raw File Name":
                            nameindex = i
                        elif file_data.row_values(8)[i] == "Calculated Amount":
                            concindex = i

                    # 未规范设置表格列名
                    if concindex == 0:
                        error_message = "未规范设置文件列名，请检查并规范设置后重新提交数据!"
                        return {"error_message": error_message}


                    # 判断一：是否设置了loq
                    if PTloq == [] or all(i is None for i in PTloq): # 第一种情况，未设置loq
                    # 判断是否在后台管理系统中设置了可接受标准

                        # 第一种情况，未设置，前端调用可接受区间模板
                        if all(i is None for i in PTrange1):
                            templates = 1
                            for i in range(nrows):
                                if "PT" in file_data.row_values(i)[nameindex]:

                                    # 第一种情况：未稀释，结果不需要除以稀释倍数
                                    if "times" not in file_data.row_values(i)[nameindex]:
                                        PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex], effectnum(file_data.row_values(i)[concindex], digits)])

                                    # 第二种情况：进行稀释，结果需要除以稀释倍数(实验号命名规则：PT-2times-1）
                                    else:
                                        dilute = int(file_data.row_values(i)[nameindex][3:4])
                                        actualconc = float(file_data.row_values(i)[concindex])*dilute
                                        PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex], effectnum(actualconc, digits)])

                        # 第二种情况，已设置，前端调用可接受标准模板
                        else:
                            templates = 2
                            for i in range(nrows):
                                if "PT" in file_data.row_values(i)[nameindex]:

                                    # 第一种情况：未稀释，结果不需要除以稀释倍数
                                    if "times" not in file_data.row_values(i)[nameindex]:
                                        # 小于range1,添加第一个可接受标准
                                        if float(file_data.row_values(i)[concindex]) < PTrange1[j]:
                                            PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex], effectnum(file_data.row_values(i)[concindex], digits), "±"+" "+str(PTstandard1[j])+" "+PTunit])
                                        # 大于range2,添加第二个可接受标准
                                        elif float(file_data.row_values(i)[concindex]) >= PTrange2[j]:
                                            PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex], effectnum(file_data.row_values(i)[concindex], digits), "±"+" "+str(PTstandard2[j])+" "+"%"])

                                    # 第二种情况：进行稀释，结果需要除以稀释倍数(实验号命名规则：PT-2times-1）
                                    else:
                                        dilute = int(file_data.row_values(i)[nameindex][3:4])
                                        actualconc = float(file_data.row_values(i)[concindex])*dilute

                                        if float(actualconc) < PTrange1[j]:
                                            PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex], effectnum(actualconc, digits), "±"+" "+str(PTstandard1[j])+" "+PTunit])
                                        elif float(actualconc) >= PTrange2[j]:
                                            PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex], effectnum(actualconc, digits), "±"+" "+str(PTstandard2[j])+" "+"%"])

                    else: # 第二种情况，设置了loq
                        if len(PTloq)!=len(PTnorm):  # loq设置不全，返回错误提示信息
                            error_message = "loq信息设置不全，请检查并规范设置后重新提交数据!"
                            return {"error_message": error_message}
                        else:
                            print(PTloq)
                            PTloq = [0 if i == 0.0 or i is None else i for i in PTloq] # 替换0.0为0
                            # 第一种情况，未设置，前端调用可接受区间模板
                            if all(i is None for i in PTrange1):
                                templates = 1
                                for i in range(nrows):
                                    if "PT" in file_data.row_values(i)[nameindex]:

                                        # 第一种情况：未稀释，结果不需要除以稀释倍数
                                        if "times" not in file_data.row_values(i)[nameindex]:
                                            # 判断检测结果是否小于等于loq
                                            try:
                                                if float(effectnum(file_data.row_values(i)[concindex], digits))<=PTloq[j]:
                                                    PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex], "<"+str(PTloq[j])])
                                                else:
                                                    PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex], effectnum(file_data.row_values(i)[concindex], digits)])
                                            except:
                                                PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex], "N/A"])

                                        # 第二种情况：进行稀释，结果需要除以稀释倍数(实验号命名规则：PT-2times-1）
                                        else:
                                            try:
                                                dilute = int(file_data.row_values(i)[nameindex][3:4])
                                                actualconc = float(file_data.row_values(i)[concindex])*dilute

                                                # 判断检测结果是否小于等于loq
                                                if float(effectnum(actualconc, digits))<=PTloq[j]:
                                                    PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex], "<"+str(PTloq[j])])
                                                else:
                                                    PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex], effectnum(actualconc, digits)])
                                            except:
                                                PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex], "N/A"])

                            # 第二种情况，已设置，前端调用可接受标准模板
                            else:
                                templates = 2
                                for i in range(nrows):
                                    if "PT" in file_data.row_values(i)[nameindex]:

                                        # 第一种情况：未稀释，结果不需要除以稀释倍数
                                        if "times" not in file_data.row_values(i)[nameindex]:

                                             # 判断检测结果是否小于等于loq
                                            try:
                                                if float(effectnum(file_data.row_values(i)[concindex], digits))<=PTloq[j]:
                                                    if float(effectnum(file_data.row_values(i)[concindex], digits)) < PTrange1[j]:
                                                        PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex], "<"+str(PTloq[j]), "±"+" "+str(PTstandard1[j])+" "+PTunit])
                                                    elif float(effectnum(file_data.row_values(i)[concindex], digits)) >= PTrange2[j]:
                                                        PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex], "<"+str(PTloq[j]), "±"+" "+str(PTstandard2[j])+" "+"%"])

                                                else:
                                                    if float(effectnum(file_data.row_values(i)[concindex], digits)) < PTrange1[j]:
                                                        PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex], effectnum(file_data.row_values(i)[concindex], digits), "±"+" "+str(PTstandard1[j])+" "+PTunit])
                                                    elif float(effectnum(file_data.row_values(i)[concindex], digits)) >= PTrange2[j]:
                                                        PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex], effectnum(file_data.row_values(i)[concindex], digits), "±"+" "+str(PTstandard2[j])+" "+"%"])
                                            except:
                                                PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex], "N/A"])

                                        # 第二种情况：进行稀释，结果需要除以稀释倍数(实验号命名规则：PT-2times-1）
                                        else:
                                            dilute = int(file_data.row_values(i)[nameindex][3:4])
                                            actualconc = float(effectnum(file_data.row_values(i)[concindex], digits))*dilute

                                            # 判断检测结果是否小于等于loq
                                            try:
                                                if float(actualconc)<=PTloq[j]:
                                                    if float(actualconc) < PTrange1[j]:
                                                        PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex], "<"+str(PTloq[j]), "±"+" "+str(PTstandard1[j])+" "+PTunit])
                                                    elif float(actualconc) >= PTrange2[j]:
                                                        PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex], "<"+str(PTloq[j]), "±"+" "+str(PTstandard2[j])+" "+"%"])

                                                else:
                                                    if float(actualconc) < PTrange1[j]:
                                                        PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex], effectnum(actualconc, digits), "±"+" "+str(PTstandard1[j])+" "+PTunit])
                                                    elif float(actualconc) >= PTrange2[j]:
                                                        PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex], effectnum(actualconc, digits), "±"+" "+str(PTstandard2[j])+" "+"%"])
                                            except:
                                                PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex], "N/A"])

            elif manufacturers == "岛津":
                # 3 读取txt
                content = []
                for line in file:
                    try:
                        content.append(line.decode("UTF-8").replace("\r\n", "").split("\t"))  # windows下
                    except:
                        content.append(line.decode("GB2312").replace("\r\n", "").split("\t")) # linux下

                nameindex = 0
                concindex = 0  # 浓度索引，岛津的数据格式决定每个化合物的浓度所在列一定是同一列
                norm = []  # 化合物列表
                norm_row = []  # 化合物所在行

                for i in range(len(content[2])):  # 第二行确定samplename和浓度所在列
                    if content[2][i] == "数据文件名":
                        nameindex = i
                    elif content[2][i] == "浓度":
                        concindex = i

                # 未规范设置表格列名
                if nameindex == 0 or concindex == []:
                    error_message = "未规范设置文件列名，请检查并规范设置后重新提交数据!"
                    return {"error_message": error_message}

                for i in range(len(content)):
                    if content[i][0] == "Name":  # 如果某一行第一列为"Name"，则该行第二列为化合物名称
                        if "-" in content[i][1]:
                            norm.append(content[i][1].split("-")[0])
                        else:
                            norm.append(content[i][1])
                        norm_row.append(i)

                # 判断是否在后台管理系统中设置了可接受标准

                # 第一种情况，未设置，前端调用可接受区间模板
                if all(i is None for i in PTrange1):
                    templates = 1
                    for j in range(len(norm)):
                        if j < len(norm)-1:  # 如果不是最后一个化合物，索引为该化合物所在行到后一个化合物所在行
                            for i in range(norm_row[j], norm_row[j+1]):
                                if "PT" in content[i][nameindex]:

                                    # 第一种情况：未稀释，结果不需要除以稀释倍数
                                    if "times" not in content[i][nameindex]:
                                        PT_dict[PTnorm[j]].append([content[i][nameindex], effectnum(content[i][concindex], digits)])

                                    # 第二种情况：进行稀释，结果需要除以稀释倍数(实验号命名规则：PT-2times-1）
                                    else:
                                        dilute = int(content[i][nameindex][3:4])
                                        actualconc = float(content[i][concindex])*dilute
                                        PT_dict[PTnorm[j]].append([content[i][nameindex], effectnum(actualconc, digits)])

                        else:
                            for i in range(norm_row[j], len(content)-1):
                                if "PT" in content[i][nameindex]:

                                    # 第一种情况：未稀释，结果不需要除以稀释倍数
                                    if "times" not in content[i][nameindex]:
                                        PT_dict[PTnorm[j]].append([content[i][nameindex], effectnum(content[i][concindex], digits)])

                                    # 第二种情况：进行稀释，结果需要除以稀释倍数(实验号命名规则：PT-2times-1）
                                    else:
                                        dilute = int(content[i][nameindex][3:4])
                                        actualconc = float(content[i][concindex])*dilute
                                        PT_dict[PTnorm[j]].append([content[i][nameindex], effectnum(actualconc, digits)])

                # 第二种情况，已设置，前端调用可接受标准模板
                else:
                    templates = 2
                    for j in range(len(norm)):
                        if j < len(norm)-1:  # 如果不是最后一个化合物，索引为该化合物所在行到后一个化合物所在行
                            for i in range(norm_row[j], norm_row[j+1]):
                                if "PT" in content[i][nameindex]:

                                    # 第一种情况：未稀释，结果不需要除以稀释倍数
                                    if "times" not in content[i][nameindex]:
                                        # 小于range1,添加第一个可接受标准
                                        if float(content[i][concindex]) < PTrange1[j]:
                                            PT_dict[PTnorm[j]].append([content[i][nameindex], effectnum(content[i][concindex], digits), "±"+" "+str(PTstandard1[j])+" "+PTunit])
                                        # 大于range2,添加第二个可接受标准
                                        elif float(content[i][concindex]) >= PTrange2[j]:
                                            PT_dict[PTnorm[j]].append([content[i][nameindex], effectnum(content[i][concindex], digits), "±"+" "+str(PTstandard2[j])+" "+"%"])

                                    # 第二种情况：进行稀释，结果需要除以稀释倍数(实验号命名规则：PT-2times-1）
                                    else:
                                        dilute = int(content[i][nameindex][3:4])
                                        actualconc = float(content[i][concindex])*dilute

                                        if float(actualconc) < PTrange1[j]:
                                            PT_dict[PTnorm[j]].append([content[i][nameindex], effectnum(actualconc, digits), "±"+" "+str(PTstandard1[j])+" "+PTunit])
                                        elif float(actualconc) >= PTrange2[j]:
                                            PT_dict[PTnorm[j]].append([content[i][nameindex], effectnum(actualconc, digits), "±"+" "+str(PTstandard2[j])+" "+"%"])

                        else:  # 如果是最后一个化合物，索引为该化合物所在行到总行数
                            for i in range(norm_row[j], len(content)-1):
                                if "PT" in content[i][nameindex]:

                                    # 第一种情况：未稀释，结果不需要除以稀释倍数
                                    if "times" not in content[i][nameindex]:
                                        # 小于range1,添加第一个可接受标准
                                        if float(content[i][concindex]) < PTrange1[j]:
                                            PT_dict[PTnorm[j]].append([content[i][nameindex], effectnum(content[i][concindex], digits), "±"+" "+str(PTstandard1[j])+" "+PTunit])
                                        # 大于range2,添加第二个可接受标准
                                        elif float(content[i][concindex]) >= PTrange2[j]:
                                            PT_dict[PTnorm[j]].append([content[i][nameindex], effectnum(content[i][concindex], digits), "±"+" "+str(PTstandard2[j])+" "+"%"])

                                    # 第二种情况：进行稀释，结果需要除以稀释倍数(实验号命名规则：PT-2times-1）
                                    else:
                                        dilute = int(content[i][nameindex][3:4])
                                        actualconc = float(content[i][concindex])*dilute

                                        if float(actualconc) < PTrange1[j]:
                                            PT_dict[PTnorm[j]].append([content[i][nameindex], effectnum(actualconc, digits), "±"+" "+str(PTstandard1[j])+" "+PTunit])
                                        elif float(actualconc) >= PTrange2[j]:
                                            PT_dict[PTnorm[j]].append([content[i][nameindex], effectnum(actualconc, digits), "±"+" "+str(PTstandard2[j])+" "+"%"])

            elif manufacturers == "AB":
                print(PTloq)
                # 获取上传的文件
                file_data = Document(file)

                # 每个表格最上方的标题内容列表，含有母离子和子离子的信息。需依此及母离子和子离子列表判断table索引
                paragraphs = []

                # 若标题长度为0或为换行等，不添加进入标题内容列表
                for p in file_data.paragraphs:
                    if len(p.text) != 0 and p.text != "\n" and len(p.text.strip()) != 0:
                        paragraphs.append(p.text.strip())

                print(paragraphs)

                # 确定table索引，母离子和子离子都与后台管理系统中设置的数值相同才证明是需要读取的定量表格
                tableindex = []
                for i in range(len(paragraphs)):
                    for j in range(len(ZP_Method_precursor_ion)):
                        if ZP_Method_precursor_ion[j] in paragraphs[i] and ZP_Method_product_ion[j] in paragraphs[i]:
                            tableindex.append(2*i+1)

                tables = file_data.tables  # 获取文件中的表格集

                # 判断一：是否设置了loq
                if PTloq == [] or all(i is None for i in PTloq): # 第一种情况，未设置loq
                    # 循环定量表格的索引
                    for k in range(len(tableindex)):
                        tablequantify = tables[tableindex[k]]  # 获取文件中的相关表格

                        # 先把表格里的所有数据取出来放进一个列表中，读取速度会比直接读表格快很多
                        cells = tablequantify._cells
                        ROWS = len(tablequantify.rows)
                        COLUMNS = len(tablequantify.columns)
                        rowdatalist = []  # 每一行的数据列表
                        rowdatagatherlist = []  # 每一行的数据列表汇总列表

                        for i in range(ROWS*COLUMNS):
                            text = cells[i].text.replace("\n", "")
                            text = text.strip()  # 去除空白符
                            if i % COLUMNS != 0 or i == 0:
                                rowdatalist.append(text)
                            else:
                                rowdatagatherlist.append(rowdatalist)
                                rowdatalist = []
                                rowdatalist.append(text)
                        rowdatagatherlist.append(rowdatalist)

                        nameindex = 0
                        concindex = 0  # 浓度索引，AB的数据格式决定每个化合物的浓度所在列一定是同一列

                        # 读取表格的第一行的单元格,判断实验号和浓度索引
                        for i in range(len(rowdatagatherlist[0])):
                            if rowdatagatherlist[0][i] == "Sample Name":
                                nameindex = i
                            elif "Calculated Conc" in rowdatagatherlist[0][i]:
                                concindex = i

                        # 浓度一般不会为第一列，若concindex == 0，则未规范设置表格列名
                        if concindex == 0:
                            error_message = "未规范设置文件列名，请检查并规范设置后重新提交数据!"
                            return {"error_message": error_message}

                        # 判断是否在后台管理系统中设置了可接受标准
                        # 第一种情况，未设置，前端调用可接受区间模板
                        if all(i is None for i in PTrange1):
                            templates = 1
                            for i in range(len(rowdatagatherlist)):
                                if "PT" in rowdatagatherlist[i][nameindex]:

                                    # 第一种情况：未稀释，结果不需要除以稀释倍数
                                    if "times" not in rowdatagatherlist[i][nameindex]:
                                        try:
                                            PT_dict[PTnorm[k]].append([rowdatagatherlist[i][nameindex], effectnum(rowdatagatherlist[i][concindex], digits)])
                                        except:
                                            PT_dict[PTnorm[k]].append([rowdatagatherlist[i][nameindex], "N/A"])

                                    # 第二种情况：进行稀释，结果需要除以稀释倍数(实验号命名规则：PT-2times-1）
                                    else:
                                        try:
                                            dilute = int(rowdatagatherlist[i][nameindex][3:4])
                                            actualconc = float(rowdatagatherlist[i][concindex])*dilute
                                            PT_dict[PTnorm[k]].append([rowdatagatherlist[i][nameindex], effectnum(actualconc, digits)])
                                        except:
                                            PT_dict[PTnorm[k]].append([rowdatagatherlist[i][nameindex], "N/A"])


                        # 第二种情况，已设置，前端调用可接受标准模板
                        else:
                            templates = 2
                            for i in range(len(rowdatagatherlist)):
                                if "PT" in rowdatagatherlist[i][nameindex]:

                                    # 第一种情况：未稀释，结果不需要除以稀释倍数
                                    if "times" not in rowdatagatherlist[i][nameindex]:
                                        if float(rowdatagatherlist[i][concindex]) < PTrange1[k]:
                                            PT_dict[PTnorm[k]].append([rowdatagatherlist[i][nameindex], effectnum(rowdatagatherlist[i][concindex], digits), "±"+" "+str(PTstandard1[k])+" "+PTunit])
                                        elif float(rowdatagatherlist[i][concindex]) >= PTrange2[k]:
                                            PT_dict[PTnorm[k]].append([rowdatagatherlist[i][nameindex], effectnum(rowdatagatherlist[i][concindex], digits), "±"+" "+str(PTstandard2[k])+" "+"%"])

                                    # 第二种情况：进行稀释，结果需要除以稀释倍数(实验号命名规则：PT-2times-1）
                                    else:
                                        dilute = int(rowdatagatherlist[i][nameindex][3:4])
                                        actualconc = float(rowdatagatherlist[i][concindex])*dilute

                                        if float(actualconc) < PTrange1[k]:
                                            PT_dict[PTnorm[k]].append([rowdatagatherlist[i][nameindex], effectnum(actualconc, digits), "±"+" "+str(PTstandard1[k])+" "+PTunit])
                                        elif float(actualconc) >= PTrange2[k]:
                                            PT_dict[PTnorm[k]].append([rowdatagatherlist[i][nameindex], effectnum(actualconc, digits), "±"+" "+str(PTstandard2[k])+" "+"%"])


                else: # 第二种情况，设置了loq
                    if len(PTloq)!=len(PTnorm):  # loq设置不全，返回错误提示信息
                        error_message = "loq信息设置不全，请检查并规范设置后重新提交数据!"
                        return {"error_message": error_message}
                    else:
                        # 循环定量表格的索引

                        PTloq = [0 if i == 0.0 else i for i in PTloq] # 替换0.0为0
                        for k in range(len(tableindex)):
                            tablequantify = tables[tableindex[k]]  # 获取文件中的相关表格

                            # 先把表格里的所有数据取出来放进一个列表中，读取速度会比直接读表格快很多
                            cells = tablequantify._cells
                            ROWS = len(tablequantify.rows)
                            COLUMNS = len(tablequantify.columns)
                            rowdatalist = []  # 每一行的数据列表
                            rowdatagatherlist = []  # 每一行的数据列表汇总列表

                            for i in range(ROWS*COLUMNS):
                                text = cells[i].text.replace("\n", "")
                                text = text.strip()  # 去除空白符
                                if i % COLUMNS != 0 or i == 0:
                                    rowdatalist.append(text)
                                else:
                                    rowdatagatherlist.append(rowdatalist)
                                    rowdatalist = []
                                    rowdatalist.append(text)
                            rowdatagatherlist.append(rowdatalist)

                            nameindex = 0
                            concindex = 0  # 浓度索引，AB的数据格式决定每个化合物的浓度所在列一定是同一列

                            # 读取表格的第一行的单元格,判断实验号和浓度索引
                            for i in range(len(rowdatagatherlist[0])):
                                if rowdatagatherlist[0][i] == "Sample Name":
                                    nameindex = i
                                elif "Calculated Conc" in rowdatagatherlist[0][i]:
                                    concindex = i

                            # 浓度一般不会为第一列，若concindex == 0，则未规范设置表格列名
                            if concindex == 0:
                                error_message = "未规范设置文件列名，请检查并规范设置后重新提交数据!"
                                return {"error_message": error_message}

                            # 判断是否在后台管理系统中设置了可接受标准
                            # 第一种情况，未设置，前端调用可接受区间模板
                            if all(i is None for i in PTrange1):
                                templates = 1
                                for i in range(len(rowdatagatherlist)):
                                    if "PT" in rowdatagatherlist[i][nameindex]:

                                        # 第一种情况：未稀释，结果不需要除以稀释倍数
                                        if "times" not in rowdatagatherlist[i][nameindex]:

                                            # 判断检测结果是否小于等于loq
                                            try:
                                                if float(effectnum(rowdatagatherlist[i][concindex], digits))<=PTloq[k]:
                                                    PT_dict[PTnorm[k]].append([rowdatagatherlist[i][nameindex], "<"+str(PTloq[k])])
                                                else:
                                                    PT_dict[PTnorm[k]].append([rowdatagatherlist[i][nameindex], effectnum(rowdatagatherlist[i][concindex], digits)])
                                            except:
                                                PT_dict[PTnorm[k]].append([rowdatagatherlist[i][nameindex], "N/A"])

                                        # 第二种情况：进行稀释，结果需要除以稀释倍数(实验号命名规则：PT-2times-1）
                                        else:
                                            try:
                                                dilute = int(rowdatagatherlist[i][nameindex][3:4])
                                                actualconc = float(rowdatagatherlist[i][concindex])*dilute

                                                # 判断检测结果是否小于等于loq
                                                if float(effectnum(actualconc, digits))<=PTloq[k]:
                                                    PT_dict[PTnorm[k]].append([rowdatagatherlist[i][nameindex], "<"+str(PTloq[k])])
                                                else:
                                                    PT_dict[PTnorm[k]].append([rowdatagatherlist[i][nameindex], effectnum(actualconc, digits)])
                                            except:
                                                PT_dict[PTnorm[k]].append([rowdatagatherlist[i][nameindex], "N/A"])

                            # 第二种情况，已设置，前端调用可接受标准模板
                            else:
                                templates = 2
                                for i in range(len(rowdatagatherlist)):
                                    if "PT" in rowdatagatherlist[i][nameindex]:

                                        # 第一种情况：未稀释，结果不需要除以稀释倍数
                                        if "times" not in rowdatagatherlist[i][nameindex]:

                                            # 判断检测结果是否小于等于loq
                                            try:
                                                if float(effectnum(rowdatagatherlist[i][concindex], digits))<=PTloq[k]:
                                                    if float(effectnum(rowdatagatherlist[i][concindex], digits)) < PTrange1[k]:
                                                        PT_dict[PTnorm[k]].append([rowdatagatherlist[i][nameindex], "<"+str(PTloq[k]), "±"+" "+str(PTstandard1[k])+" "+PTunit])
                                                    elif float(effectnum(rowdatagatherlist[i][concindex], digits)) >= PTrange2[k]:
                                                        PT_dict[PTnorm[k]].append([rowdatagatherlist[i][nameindex], "<"+str(PTloq[k]), "±"+" "+str(PTstandard2[k])+" "+"%"])

                                                else:
                                                    if float(effectnum(rowdatagatherlist[i][concindex], digits)) < PTrange1[k]:
                                                        PT_dict[PTnorm[k]].append([rowdatagatherlist[i][nameindex], effectnum(rowdatagatherlist[i][concindex], digits), "±"+" "+str(PTstandard1[k])+" "+PTunit])
                                                    elif float(effectnum(rowdatagatherlist[i][concindex], digits)) >= PTrange2[k]:
                                                        PT_dict[PTnorm[k]].append([rowdatagatherlist[i][nameindex], effectnum(rowdatagatherlist[i][concindex], digits), "±"+" "+str(PTstandard2[k])+" "+"%"])
                                            except:
                                                PT_dict[PTnorm[k]].append([rowdatagatherlist[i][nameindex], "N/A"])


                                        # 第二种情况：进行稀释，结果需要除以稀释倍数(实验号命名规则：PT-2times-1）
                                        else:
                                            dilute = int(rowdatagatherlist[i][nameindex][3:4])
                                            actualconc = float(effectnum(rowdatagatherlist[i][concindex], digits))*dilute

                                            # 判断检测结果是否小于等于loq
                                            try:
                                                if float(actualconc)<=PTloq[k]:
                                                    if float(actualconc) < PTrange1[k]:
                                                        PT_dict[PTnorm[k]].append([rowdatagatherlist[i][nameindex], "<"+str(PTloq[k]), "±"+" "+str(PTstandard1[k])+" "+PTunit])
                                                    elif float(actualconc) >= PTrange2[k]:
                                                        PT_dict[PTnorm[k]].append([rowdatagatherlist[i][nameindex], "<"+str(PTloq[k]), "±"+" "+str(PTstandard2[k])+" "+"%"])

                                                else:
                                                    if float(actualconc) < PTrange1[k]:
                                                        PT_dict[PTnorm[k]].append([rowdatagatherlist[i][nameindex], effectnum(actualconc, digits), "±"+" "+str(PTstandard1[k])+" "+PTunit])
                                                    elif float(actualconc) >= PTrange2[k]:
                                                        PT_dict[PTnorm[k]].append([rowdatagatherlist[i][nameindex], effectnum(actualconc, digits), "±"+" "+str(PTstandard2[k])+" "+"%"])
                                            except:
                                                PT_dict[PTnorm[k]].append([rowdatagatherlist[i][nameindex], "N/A"])

        elif platform == "液相":
            if manufacturers == "Agilent":

                # .xlsx格式
                data = xlrd.open_workbook(filename=None, file_contents=file.read())  # 读取表格
                file_data = data.sheets()[0]
                nrows = file_data.nrows
                ncols = file_data.ncols

                norm = []  # 化合物列表
                norm_row = []  # 化合物所在行
                for j in range(nrows):
                    # 如果某一行的第一个元素为“化合物”，则添加第三个元素进入化合物列表
                    if file_data.row_values(j)[0] == "化合物:":
                        norm.append(file_data.row_values(j)[2])
                        norm_row.append(j)

                nameindex = 0
                concindex = 0

                # 第一个化合物表格确定samplename和浓度所在列，norm_row[0]为第一个化合物所在行，+1是该化合物表格位于该化合物所在行的下一行
                for i in range(len(file_data.row_values(norm_row[0]+1))):
                    if "样品名称" in file_data.row_values(norm_row[0]+1)[i]:
                        nameindex = i
                    elif "含量" in file_data.row_values(norm_row[0]+1)[i]:
                        concindex = i

                # 未规范设置表格列名
                if nameindex == 0 or concindex == []:
                    error_message = "未规范设置文件列名，请检查并规范设置后重新提交数据!"
                    return {"error_message": error_message}

                # 判断是否在后台管理系统中设置了可接受标准

                # 第一种情况，未设置，前端调用可接受区间模板
                if all(i is None for i in PTrange1):
                    templates = 1
                    for j in range(len(norm)):
                        if j < len(norm)-1:  # 如果不是最后一个化合物，索引为该化合物所在行到后一个化合物所在行
                            for i in range(norm_row[j], norm_row[j+1]):
                                if "PT" in file_data.row_values(i)[nameindex]:

                                    # 第一种情况：未稀释，结果不需要除以稀释倍数
                                    if "times" not in file_data.row_values(i)[nameindex]:
                                        PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex], effectnum(file_data.row_values(i)[concindex], digits)])

                                    # 第二种情况：进行稀释，结果需要除以稀释倍数(实验号命名规则：PT-2times-1）
                                    else:
                                        dilute = int(file_data.row_values(i)[nameindex][3:4])
                                        actualconc = float(file_data.row_values(i)[concindex])*dilute
                                        PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex], effectnum(actualconc, digits)])

                        else:
                            for i in range(norm_row[j], nrows):
                                if "PT" in file_data.row_values(i)[nameindex]:

                                    # 第一种情况：未稀释，结果不需要除以稀释倍数
                                    if "times" not in file_data.row_values(i)[nameindex]:
                                        PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex], effectnum(file_data.row_values(i)[concindex], digits)])

                                    # 第二种情况：进行稀释，结果需要除以稀释倍数(实验号命名规则：PT-2times-1）
                                    else:
                                        dilute = int(file_data.row_values(i)[nameindex][3:4])
                                        actualconc = float(file_data.row_values(i)[concindex])*dilute
                                        PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex], effectnum(actualconc, digits)])

                # 第二种情况，已设置，前端调用可接受标准模板
                else:
                    templates = 2
                    for j in range(len(norm)):
                        if j < len(norm)-1:  # 如果不是最后一个化合物，索引为该化合物所在行到后一个化合物所在行
                            for i in range(norm_row[j], norm_row[j+1]):
                                if "PT" in file_data.row_values(i)[nameindex]:

                                    # 第一种情况：未稀释，结果不需要除以稀释倍数
                                    if "times" not in file_data.row_values(i)[nameindex]:
                                        # 小于range1,添加第一个可接受标准
                                        if float(file_data.row_values(i)[concindex]) < PTrange1[j]:
                                            PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex], effectnum(file_data.row_values(i)[concindex], digits), "±"+" "+str(PTstandard1[j])+" "+PTunit])
                                        # 大于range2,添加第二个可接受标准
                                        elif float(file_data.row_values(i)[concindex]) >= PTrange2[j]:
                                            PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex], effectnum(file_data.row_values(i)[concindex], digits), "±"+" "+str(PTstandard2[j])+" "+"%"])

                                    # 第二种情况：进行稀释，结果需要除以稀释倍数(实验号命名规则：PT-2times-1）
                                    else:
                                        dilute = int(file_data.row_values(i)[nameindex][3:4])
                                        actualconc = float(file_data.row_values(i)[concindex])*dilute

                                        if float(actualconc) < PTrange1[j]:
                                            PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex], effectnum(actualconc, digits), "±"+" "+str(PTstandard1[0])+" "+PTunit])
                                        elif float(actualconc) >= PTrange2[j]:
                                            PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex], effectnum(actualconc, digits), "±" + " " + str(PTstandard2[0]) + " " + "%"])
                        
                        else:  # 如果是最后一个化合物，索引为该化合物所在行到总行数
                            for i in range(norm_row[j], nrows):
                                if "PT" in file_data.row_values(i)[nameindex]:
                                    # 小于range1,添加第一个可接受标准
                                    if float(effectnum(file_data.row_values(i)[concindex], digits)) < PTrange1[j]:
                                        PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex], effectnum(file_data.row_values(i)[concindex], digits), "±"+" "+str(PTstandard1[j])+" "+PTunit])
                                    # 大于range2,添加第二个可接受标准
                                    elif float(effectnum(file_data.row_values(i)[concindex], digits)) >= PTrange2[j]:
                                        PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex], effectnum(file_data.row_values(i)[concindex], digits), "±"+" "+str(PTstandard2[j])+" "+"%"])

        elif platform == "ICP-MS":
            if manufacturers == "Agilent":
                data = xlrd.open_workbook(filename=None, file_contents=file.read())  # 读取表格
                file_data = data.sheets()[0]
                nrows = file_data.nrows
                ncols = file_data.ncols

                # 从第一行确定化合物名称
                norm = []
                for j in range(ncols):
                    for i in PTnorm:
                        if i in file_data.row_values(0)[j]:
                            norm.append(i)

                # 从第二行确定实验号（Sample Name）的索引和化合物浓度索引
                nameindex = 0  # 实验号索引
                concindex = []  # 浓度索引
                for j in range(ncols):
                    if file_data.row_values(1)[j] == "样品名称":
                        nameindex = j
                    elif file_data.row_values(1)[j] == "浓度 [ ppm ]" or file_data.row_values(1)[j] == "浓度 [ ppb ]":
                        concindex.append(j)

                # 未规范设置表格列名
                if nameindex == 0 or concindex == []:
                    error_message = "未规范设置文件列名，请检查并规范设置后重新提交数据!"
                    return {"error_message": error_message}

                # 判断是否在后台管理系统中设置了可接受标准

                # 第一种情况，未设置，前端调用可接受区间模板
                if all(i is None for i in PTrange1):
                    templates = 1
                    for j in range(len(concindex)):
                        for i in range(2, nrows):  # 循环原始数据中的每一行
                            if "PT" in file_data.row_values(i)[nameindex]:

                                # 第一种情况：未稀释，结果不需要除以稀释倍数
                                if "B" not in file_data.row_values(i)[nameindex]:
                                    PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex], effectnum(file_data.row_values(i)[concindex[j]], digits)])
                                

                                # 第二种情况：进行稀释，结果需要除以稀释倍数(实验号命名规则：PT-2times-1）
                                else:
                                    dilute = int(file_data.row_values(i)[nameindex][3:4])
                                    actualconc = float(file_data.row_values(i)[concindex[j]])*dilute
                                    PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex], effectnum(actualconc, digits)])

                # 第二种情况，已设置，前端调用可接受标准模板
                else:
                    templates = 2
                    # 匹配原始数据中与PT相关(实验号前含有"PT")的行
                    for j in range(len(concindex)):
                        for i in range(2, nrows):  # 循环原始数据中的每一行
                            if "PT" in file_data.row_values(i)[nameindex]:

                                # 第一种情况：未稀释，结果不需要除以稀释倍数
                                    if "B" not in file_data.row_values(i)[nameindex]:
                                        # 小于range1,添加第一个可接受标准
                                        if float(file_data.row_values(i)[concindex]) < PTrange1[j]:
                                            PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex], effectnum(file_data.row_values(i)[concindex], digits), "±"+" "+str(PTstandard1[j])+" "+PTunit])
                                        # 大于range2,添加第二个可接受标准
                                        elif float(file_data.row_values(i)[concindex]) >= PTrange2[j]:
                                            PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex], effectnum(file_data.row_values(i)[concindex], digits), "±"+" "+str(PTstandard2[j])+" "+"%"])

                                    # 第二种情况：进行稀释，结果需要除以稀释倍数(实验号命名规则：PT-2times-1）
                                    else:
                                        dilute = int(file_data.row_values(i)[nameindex][3:4])
                                        actualconc = float(file_data.row_values(i)[concindex])*dilute

                                        if float(actualconc) < PTrange1[j]:
                                            PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex], effectnum(actualconc, digits), "±"+" "+str(PTstandard1[j])+" "+PTunit])
                                        elif float(actualconc) >= PTrange2[j]:
                                            PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex], effectnum(actualconc, digits), "±"+" "+str(PTstandard2[j])+" "+"%"])


            PT_num = len(PT_dict[PTnorm[0]])

        # 判断每个指标有几个样本
        PT_num = len(PT_dict[PTnorm[0]])
        print(PT_dict)

    return {"PT_dict": PT_dict, "PT_num": PT_num, "PTunit": PTunit, "templates": templates, "project": project}


def PT_25OHD_fileread(files, Detectionplatform, project, platform, manufacturers, digits, ZP_Method_precursor_ion, ZP_Method_product_ion, normAB, Number_of_compounds):

    # 第一步 后台数据抓取（待测物质,可接受标准，单位）
    zqd = Special.objects.get(project=project)
    pt_special = PTspecial.objects.get(special=zqd)
    pt_accept = PTspecialaccept.objects.filter(pTspecial=pt_special)
    PTnorm = []  # 待测物质列表
    PTrange1 = []  # 可接受标准一适用范围，与待测物质列表一一对应
    PTstandard1 = []  # 可接受标准一
    PTrange2 = []  # 可接受标准二适用范围，与待测物质列表一一对应
    PTstandard2 = []  # 可接受标准二
    PTunit = Special.objects.get(project=project).unit  # 单位

    for i in pt_accept:
        PTnorm.append(i.norm)
        PTrange1.append(i.range1)
        PTstandard1.append(i.accept1)
        PTrange2.append(i.range2)
        PTstandard2.append(i.accept2)

    # 如果没在后台管理系统中设置化合物名称直接返回并提示
    if all(i is None for i in PTnorm):
        error_message = "尚未在后台管理系统中设置PT的化合物名称,请设置后重新提交数据!"
        return {"error_message": error_message}

    # Python判断列表中是否为空，包括None
    # if all(i is None for i in PTrange1):
    #     error_message="尚未在后台管理系统中设置PT的可接受标准,请设置后重新提交数据!"
    #     return {"error_message":error_message}

    # AB厂家,未在后台管理系统里规范设置定量离子对数值,直接返回并提示
    if manufacturers == "AB":
        if len(normAB) != Number_of_compounds:
            error_message = "未在后台管理系统里规范设置定量离子对数值，请检查并规范设置后重新提交数据!"
            return {"error_message": error_message}

    #  第二步:开始文件读取

    # 头部定义相关需要提取生成的结果
    PT_dict = {}

    # 判断是否在后台管理系统中设置了可接受标准

    # 第一种情况，未设置，前端调用可接受区间模板
    if all(i is None for i in PTrange1):
        templates = 1

    # 第二种情况，已设置，前端调用可接受标准模板
    else:
        templates = 2

    # 各仪器平台及各仪器厂家数据读取
    for file in files:
        if platform == "液质":
            if manufacturers == "AB":
                # 获取上传的文件
                file_data = Document(file)

                # 每个表格最上方的标题内容列表，含有母离子和子离子的信息。需依此及母离子和子离子列表判断table索引
                paragraphs = []

                # 若标题长度为0或为换行等，不添加进入标题内容列表
                for p in file_data.paragraphs:
                    if len(p.text) != 0 and p.text != "\n" and len(p.text.strip()) != 0:
                        paragraphs.append(p.text.strip())

                print(paragraphs)

                # 确定table索引，母离子和子离子都与后台管理系统中设置的数值相同才证明是需要读取的定量表格
                tableindex = []
                for i in range(len(paragraphs)):
                    for j in range(len(ZP_Method_precursor_ion)):
                        if ZP_Method_precursor_ion[j] in paragraphs[i] and ZP_Method_product_ion[j] in paragraphs[i]:
                            tableindex.append(2*i+1)

                print(tableindex)

                tables = file_data.tables  # 获取文件中的表格集

                # 循环定量表格的索引
                for k in range(len(tableindex)):
                    PT_dict[normAB[k]] = []
                    tablequantify = tables[tableindex[k]]  # 获取文件中的相关表格
                    # nameindex = 0
                    # conindex = 0

                    # 先把表格里的所有数据取出来放进一个列表中，读取速度会比直接读表格快很多
                    cells = tablequantify._cells
                    ROWS = len(tablequantify.rows)
                    COLUMNS = len(tablequantify.columns)
                    rowdatalist = []  # 每一行的数据列表
                    rowdatagatherlist = []  # 每一行的数据列表汇总列表

                    for i in range(ROWS*COLUMNS):
                        text = cells[i].text.replace("\n", "")
                        text = text.strip()  # 去除空白符
                        if i % COLUMNS != 0 or i == 0:
                            rowdatalist.append(text)
                        else:
                            rowdatagatherlist.append(rowdatalist)
                            rowdatalist = []
                            rowdatalist.append(text)
                    rowdatagatherlist.append(rowdatalist)

                    nameindex = 0
                    concindex = 0  # 浓度索引，AB的数据格式决定每个化合物的浓度所在列一定是同一列

                    # 读取表格的第一行的单元格,判断实验号和浓度索引
                    for i in range(len(rowdatagatherlist[0])):
                        if rowdatagatherlist[0][i] == "Sample Name":
                            nameindex = i
                        elif "Calculated Conc" in rowdatagatherlist[0][i]:
                            concindex = i
                    
                    # 未规范设置表格列名
                    if concindex == 0:
                        error_message = "未规范设置文件列名，请检查并规范设置后重新提交数据!"
                        return {"error_message": error_message}

                    for i in range(len(rowdatagatherlist)):
                        if "PT" in rowdatagatherlist[i][nameindex]:

                            # 第一种情况：未稀释，结果不需要除以稀释倍数
                            if "times" not in rowdatagatherlist[i][nameindex]:
                                if "<" in rowdatagatherlist[i][concindex]:
                                    PT_dict[normAB[k]].append([rowdatagatherlist[i][nameindex], rowdatagatherlist[i][concindex].split("<")[1]])
                                else:
                                    PT_dict[normAB[k]].append([rowdatagatherlist[i][nameindex], rowdatagatherlist[i][concindex]])

                            # 第二种情况：进行稀释，结果需要除以稀释倍数(实验号命名规则：PT-2times-1）
                            else:
                                dilute = int(rowdatagatherlist[i][nameindex][3:4])
                                actualconc = float(rowdatagatherlist[i][concindex])*dilute
                                PT_dict[normAB[k]].append([rowdatagatherlist[i][nameindex], actualconc])

            elif manufacturers == "Agilent":
                # 1 读取csv文件（Agilent）
                file.seek(0)  # https://www.jianshu.com/p/0d15ed85df2b
                file_data = file.read().decode('utf-8')
                lines = file_data.split('\r\n')
                for i in range(len(lines)):
                    if len(lines[i]) != 0:
                        # 以逗号分隔字符串,但忽略双引号内的逗号
                        lines[i] = re.split(r',\s*(?![^"]*\"\,)', lines[i])
                        # lines[i]=lines[i].split(',') # 按逗号分隔后把每一行都变成一个列表
                    else:
                        lines[i] = re.split(r',\s*(?![^"]*\"\,)', lines[i])
                        del lines[i]  # 最后一行如为空行，则删除该元素

                # 从第一行确定化合物名称(含有"-Q Results"),并添加进入化合物列表
                norm = []  # 化合物列表
                for j in range(len(lines[0])):  # 从第一行开始
                    if "-Q Results" in lines[0][j]:
                        # 若原始字符串中含有','，切割完后首位会多出一个'"',需去除
                        if lines[0][j].split("-Q")[0][0] != '"':
                            norm.append(lines[0][j].split("-Q")[0])
                        else:
                            norm.append(lines[0][j].split("-Q")[0][1:])

                # 从第二行确定实验号（Sample Name）,浓度（Exp. Conc.）的索引
                nameindex = 0  # 实验号索引
                concindex = []  # 浓度索引列表（可能不止一个化合物，因此需要把索引放在一个列表里）
                for j in range(len(lines[1])):  # 从第二行开始
                    if lines[1][j] == "Name":
                        nameindex = j
                    elif lines[1][j] == "Final Conc." or lines[1][j] == "Calc. Conc.":
                        concindex.append(j)
                
                # 未规范设置表格列名
                if concindex == []:
                    error_message = "未规范设置文件列名，请检查并规范设置后重新提交数据!"
                    return {"error_message": error_message}

                 # 匹配原始数据中与PT相关(实验号前含有"PT")的行
                for j in range(len(concindex)):
                    PT_dict[normAB[j]]=[]
                    for i in range(len(lines)):  # 循环原始数据中的每一行
                        if "PT" in lines[i][nameindex]:
                            # 第一种情况：未稀释，结果不需要除以稀释倍数
                            if "times" not in lines[i][nameindex]:
                                PT_dict[normAB[j]].append([lines[i][nameindex], lines[i][concindex[j]]])

                            # 第二种情况：进行稀释，结果需要除以稀释倍数(实验号命名规则：PT-2times-1）
                            else:
                                dilute = int(lines[i][nameindex][3:4])
                                actualconc = float(lines[i][concindex[j]])*dilute
                                PT_dict[normAB[j]].append([lines[i][nameindex], actualconc])

            # 判断每个指标有几个样本
            PT_num = len(PT_dict[normAB[0]])

            # Agilent需将PT_dict中key的顺序颠倒，value值不变
            if manufacturers=="Agilent":
                try:
                    PT_dict_trans={}
                    PT_dict_trans["25-OH-D2"]=PT_dict[normAB[1]]
                    PT_dict_trans["25-OH-D3"]=PT_dict[normAB[0]]
                    PT_dict =  PT_dict_trans
                except:
                    error_message = "未规范设置仪器条件（D2需排在D3前），请检查并规范设置后重新提交数据!"
                    return {"error_message": error_message}

        print(PT_dict)

        # 判断是血清25OHD还是末梢血25OHD
        if 'PTAD-25OHD2' in normAB:
            # 25OHD项目需计算总D，目前D2和D3的实验号顺序必须一致，否则计算会出错

            # 判断后台管理系统中设置了几个化合物
            # 第一种情况：只设置了25OHD，前端只显示一个化合物
            if len(PTnorm) == 1:
                PT_dict2 = {"PTAD-25OHD": []}
                for index, value in enumerate(PT_dict['PTAD-25OHD2']):
                    middle_list = []  # 每个实验号的列表

                    value_D2 = PT_dict['PTAD-25OHD2'][index][1]  # 每个实验号列表的第二位为结果
                    value_D3 = PT_dict['PTAD-25OHD3'][index][1]

                    # D2小于定量限(5.33)的值不需要加和
                    if float(value_D2) < 5.33:  # 定量限5.33，目前先写死
                        value_D2 = "<5.33"
                        value_D3 = effectnum(value_D3, digits)
                        value_D = effectnum(value_D3, digits)

                    else:
                        # 先再取有效位数，再加和，再取有效位数
                        value_D2 = effectnum(value_D2, digits)
                        value_D3 = effectnum(value_D3, digits)
                        value_D = effectnum((float(value_D2)+float(value_D3)), digits)

                    middle_list.append(PT_dict['PTAD-25OHD2'][index][0])  # 添加实验号
                    middle_list.append(value_D2)  # 添加D2结果
                    middle_list.append(value_D3)  # 添加D3结果
                    middle_list.append(value_D)  # 添加总D结果

                    # 判断是否设置了可接受标准
                    if templates == 1:
                        pass
                    else:
                        # 添加可接受标准
                        if float(value_D) < PTrange1[0]:
                            middle_list.append("±"+" "+str(PTstandard1[0])+" "+PTunit)
                        elif float(value_D) >= PTrange2[0]:
                            middle_list.append("±"+" "+str(PTstandard2[0])+" "+"%")
                    PT_dict2['PTAD-25OHD'].append(middle_list)

            # 第二种情况：同时设置了D2,D3和总D，前端显示三个化合物
            elif len(PTnorm) == 3:
                pass

            PT_dict = PT_dict2


        else:
            # 25OHD项目需计算总D，目前D2和D3的实验号顺序必须一致，否则计算会出错

            # 判断后台管理系统中设置了几个化合物
            # 第一种情况：只设置了25OHD，前端只显示一个化合物
            if len(PTnorm) == 1:
                PT_dict2 = {"25-OH-D": []}
                for index, value in enumerate(PT_dict['25-OH-D2']):
                    middle_list = []  # 每个实验号的列表

                    value_D2 = PT_dict['25-OH-D2'][index][1]  # 每个实验号列表的第二位为结果
                    value_D3 = PT_dict['25-OH-D3'][index][1]

                    # D2小于定量限(5.33)的值不需要加和
                    if float(value_D2) < 5.33:  # 定量限5.33，目前先写死
                        value_D2 = "<5.33"
                        value_D3 = effectnum(value_D3, digits)
                        value_D = effectnum(value_D3, digits)

                    else:
                        # 先再取有效位数，再加和，再取有效位数
                        value_D2 = effectnum(value_D2, digits)
                        value_D3 = effectnum(value_D3, digits)
                        value_D = effectnum((float(value_D2)+float(value_D3)), digits)

                    middle_list.append(PT_dict['25-OH-D2'][index][0])  # 添加实验号
                    middle_list.append(value_D2)  # 添加D2结果
                    middle_list.append(value_D3)  # 添加D3结果
                    middle_list.append(value_D)  # 添加总D结果

                    # 判断是否设置了可接受标准
                    if templates == 1:
                        pass
                    else:
                        # 添加可接受标准
                        if float(value_D) < PTrange1[0]:
                            middle_list.append("±"+" "+str(PTstandard1[0])+" "+PTunit)
                        elif float(value_D) >= PTrange2[0]:
                            middle_list.append("±"+" "+str(PTstandard2[0])+" "+"%")
                    PT_dict2['25-OH-D'].append(middle_list)

            # 第二种情况：同时设置了D2,D3和总D，前端显示三个化合物
            elif len(PTnorm) == 3:
                pass

            PT_dict = PT_dict2

        return {"PT_dict": PT_dict, "PT_num": PT_num, "PTunit": PTunit, "templates": templates, "project": project}


def Recyclefileread(files, Detectionplatform, project, platform, manufacturers, Unit, digits, ZP_Method_precursor_ion, ZP_Method_product_ion, normAB, Number_of_compounds):

    # 第一步:后台数据抓取（回收率上下限，最大允许CV）
    id1 = Special.objects.get(project=project).id
    id2 = Recyclespecial.objects.get(special_id=id1).id

    # 优先查找特殊参数设置里是否有数据，如有就调用，没有则调用通用性参数设置里的数据
    if Recyclespecialmethod.objects.filter(recyclespecial=id2):
        lowvalue = Recyclespecialmethod.objects.get(recyclespecial=id2).lowvalue  # 回收率下限
        upvalue = Recyclespecialmethod.objects.get(recyclespecial=id2).upvalue  # 回收率上限

    else:
        general = General.objects.get(name="通用性项目")
        recycle_general = Recyclegeneral.objects.get(general=general)
        lowvalue = Recyclegeneralmethod.objects.get(recyclegeneral=recycle_general).lowvalue  # 回收率下限
        upvalue = Recyclegeneralmethod.objects.get(recyclegeneral=recycle_general).upvalue  # 回收率上限

    # AB厂家,未在后台管理系统里规范设置定量离子对数值,直接返回并提示
    if manufacturers == "AB":
        if len(normAB) != Number_of_compounds:
            error_message = "未在后台管理系统里规范设置定量离子对数值，请检查并规范设置后重新提交数据!"
            return {"error_message": error_message}

    #  第二步:开始文件读取

    '''
    数据读取完毕后,需要生成一个字典Recycle_enddict,数据格式如下：
        {'待测物质1':{'sam1': [1.08, 0.44, 0.7, 1.12, 0.72, 0.76, 0.74, 1.01, 0.96, 1.23, 1.28, 0.85], 
        'sam2': [7.57, 8.33, 9.05, 13.69, 14.45, 13.36, 20.06, 19.24, 19.97, 30.72, 31.01, 30.64], 
        'sam3': [6.3, 6.48, 6.4, 9.68, 10.13, 10.52,12.24, 13.37, 14.89, 19.76, 20.77, 20.63]},
        '待测物质2':{'sam1': [1.08, 0.44, 0.7, 1.12, 0.72, 0.76, 0.74, 1.01, 0.96, 1.23, 1.28, 0.85], 
        'sam2': [7.57, 8.33, 9.05, 13.69, 14.45, 13.36, 20.06, 19.24, 19.97, 30.72, 31.01, 30.64], 
        'sam3': [6.3, 6.48, 6.4, 9.68, 10.13, 10.52, 12.24, 13.37, 14.89, 19.76, 20.77, 20.63]} }
    '''

    # 1 抓取加标回收率加标数据

    # 头部定义相关需要提取生成的结果
    Rec_dict = {}  # 加标数据字典
    Reclist = ["Rec-1", "Rec-2", "Rec-3"]  # 加标后样本实验号前缀

    for file in files:
        print(file.name)
        if file.name == "加标回收率-加标数据.csv":
            # 读取csv文件
            file.seek(0)  # 此网址查找到的答案:https://www.jianshu.com/p/0d15ed85df2b
            # print(chardet.detect(file.read()))["encoding"]
            try:
                file_data = file.read().decode('utf-8')
            except:
                file_data = file.read().decode('GB2312')
            lines = file_data.split('\r\n')
            for i in range(len(lines)):
                if len(lines[i]) != 0:
                    # 以逗号分隔字符串,但忽略双引号内的逗号
                    lines[i] = re.split(r',\s*(?![^"]*\"\,)', lines[i])
                    # lines[i]=lines[i].split(',') # 按逗号分隔后把每一行都变成一个列表
                else:
                    lines[i] = re.split(r',\s*(?![^"]*\"\,)', lines[i])
                    del lines[i]  # 最后一行如为空行，则删除该元素
            print("56789")
            print(lines)
            for i in range(len(Reclist)):
                middle_list = []  # 每个本底的加标数据列表
                print(Reclist[i])
                for j in range(len(lines)):
                    print("56789")
                    print(lines[j])
                    if Reclist[i] in lines[j][0]:
                        middle_list.extend([lines[j][1], lines[j][2], lines[j][3]])

                Rec_dict[Reclist[i]] = middle_list
            print("12345")
            print(Rec_dict)

    # 2 抓取加标回收率原始数据

    # 头部定义相关需要提取生成的结果
    Recycle_enddict_show = {}  # 加班回收率最终字典(展示数据用,如用户上传了加标数据的文件,需要添加加标数据)
    # 加班回收率最终字典(保存数据用,如用户上传了加标数据的文件,不需要添加加标数据,在验证界面点击保存按钮后需要用到此字典)
    Recycle_enddict_savedata = {}

    for file in files:
        print(file.name)
        if file.name != "加标回收率-加标数据.csv":
            if platform == "液质":
                if manufacturers == "Agilent":
                    # 1 读取csv文件（Agilent）
                    # 此网址查找到的答案:https://www.jianshu.com/p/0d15ed85df2b
                    csv_file = file.seek(0)
                    file_data = file.read().decode('utf-8')
                    lines = file_data.split('\r\n')
                    for i in range(len(lines)):
                        if len(lines[i]) != 0:
                            # 以逗号分隔字符串,但忽略双引号内的逗号
                            lines[i] = re.split(r',\s*(?![^"]*\"\,)', lines[i])
                            # lines[i]=lines[i].split(',') # 按逗号分隔后把每一行都变成一个列表
                        else:
                            lines[i] = re.split(r',\s*(?![^"]*\"\,)', lines[i])
                            del lines[i]  # 最后一行如为空行，则删除该元素

                    # 药物浓度平台的数据导出格式较不同
                    if Detectionplatform=="治疗药物检测平台":
                        norm = []  # 化合物列表
                        for j in range(len(lines[0])):  # 从第一行开始
                            if " Results" in lines[0][j]:
                                # 若原始字符串中含有','，切割完后首位会多出一个'"',需去除
                                if lines[0][j].split(" Results")[0][0] != '"':
                                    norm.append(lines[0][j].split(" Results")[0])
                                else:
                                    norm.append(lines[0][j].split(" Results")[0][1:])

                        # 定性指标的索引
                        dxindex = []
                        for i in range(len(norm)):
                            if "2" in norm[i]:
                                dxindex.append(i)

                        # 删除norm中的定性指标和定量指标的"1"
                        norm=[norm[i] for i in range(len(norm)) if("2" not in norm[i])]
                        norm=[norm[i].split("1")[0] for i in range(len(norm)) if("1" in norm[i])]
                        
                        # 从第二行确定实验号（Sample Name）,浓度（Exp. Conc.）的索引
                        nameindex = 0  # 实验号索引
                        concindex = []  # 浓度索引列表（可能不止一个化合物，因此需要把索引放在一个列表里）
                        for j in range(len(lines[1])):  # 从第二行开始
                            if lines[1][j] == "Name":
                                nameindex = j
                            elif lines[1][j] == "Final Conc.":
                                concindex.append(j)
                        
                        # 删除concindex中的dxindex
                        concindex=[concindex[i] for i in range(len(concindex)) if(i not in dxindex)]
                        print(concindex)

                    else:
                        # 从第一行确定化合物名称(含有"-Q Results"),并添加进入化合物列表
                        norm = []  # 化合物列表
                        for j in range(len(lines[0])):  # 从第一行开始
                            if "-Q Results" in lines[0][j]:
                                # 若原始字符串中含有','，切割完后首位会多出一个'"',需去除
                                if lines[0][j].split("-Q")[0][0] != '"':
                                    norm.append(lines[0][j].split("-Q")[0])
                                else:
                                    norm.append(lines[0][j].split("-Q")[0][1:])

                        # 从第二行确定实验号（Sample Name）,浓度（Exp. Conc.）的索引
                        nameindex = 0  # 实验号索引
                        concindex = []  # 浓度索引列表（可能不止一个化合物，因此需要把索引放在一个列表里）
                        for j in range(len(lines[1])):  # 从第二行开始
                            if lines[1][j] == "Name":
                                nameindex = j
                            elif lines[1][j] == "Calc. Conc.":
                                concindex.append(j)

                    # 未规范设置表格列名
                    if concindex == 0:
                        error_message = "未规范设置文件列名，请检查并规范设置后重新提交数据!"
                        return {"error_message": error_message}

                    # 确定本底数(含有"RecB")
                    RecBlist = []  # 本底列表,长度/3即为本底数
                    for i in range(len(lines)):  # 循环原始数据中的每一行,并避免重复添加
                        # 加标回收率本底实验号命名“RecB-1-1”
                        if "RecB" in lines[i][nameindex][0:6] and lines[i][nameindex][0:6] not in RecBlist:
                            RecBlist.append(lines[i][nameindex][0:6])

                    # 匹配原始数据中与加标回收相关的行
                    for k in range(len(norm)):
                        norm_dict = {}  # 每个化合物数据字典
                        for j in range(len(RecBlist)):  # 本底列表
                            RecB_conc = []  # 本底浓度列表
                            low = []  # 本底加标后低浓度列表
                            median = []  # 本底加标后中浓度列表
                            high = []  # 本底加标后高浓度列表
                            for i in range(len(lines)):  # 循环原始数据中的每一行
                                if RecBlist[j] in lines[i][nameindex]:
                                    RecB_conc.append(
                                        effectnum(lines[i][concindex[k]], digits))
                                elif "L" in lines[i][nameindex] and Reclist[j] in lines[i][nameindex]:
                                    low.append(
                                        effectnum(lines[i][concindex[k]], digits))
                                elif "M" in lines[i][nameindex] and Reclist[j] in lines[i][nameindex]:
                                    median.append(
                                        effectnum(lines[i][concindex[k]], digits))
                                elif "H" in lines[i][nameindex] and Reclist[j] in lines[i][nameindex]:
                                    high.append(
                                        effectnum(lines[i][concindex[k]], digits))

                            norm_dict[RecBlist[j]] = []

                            norm_dict[RecBlist[j]].extend(RecB_conc)
                            norm_dict[RecBlist[j]].extend(low)
                            norm_dict[RecBlist[j]].extend(median)
                            norm_dict[RecBlist[j]].extend(high)
                            if Rec_dict != {}:
                                norm_dict[RecBlist[j]].extend(Rec_dict[Reclist[j]])
                            else:
                                norm_dict[RecBlist[j]].extend(["", "", "", "", "", "", "", "", ""])

                        Recycle_enddict_show[norm[k]] = norm_dict

                    print(Recycle_enddict_show)

                elif manufacturers == "Waters":
                    # 内标标识
                    ISlist = ["D3", "D4", "D5", "D6", "D7", "D8", "dx", "d4", "d8"]

                    data = xlrd.open_workbook(filename=None, file_contents=file.read())  # 读取表格
                    file_data = data.sheets()[0]
                    nrows = file_data.nrows
                    ncols = file_data.ncols

                    norm = []  # 化合物列表
                    Compound_row = []  # 含有“Compound”关键词的所在行(包含内标)
                    norm_row = []  # 实际化合物所在行(不包含内标)
                    for i in range(nrows):
                        if "Compound" in str(file_data.row_values(i)[0]) and ":" in str(file_data.row_values(i)[0]):
                            Compound_row.append(i)

                        # 判断是否含有内标标识
                        if all(j not in str(file_data.row_values(i)[0]) for j in ISlist):
                            # 如果某一行第一列含有关键词"Compound"，则该行中含有化合物名称，化合物名称在：后
                            if "Compound" in str(file_data.row_values(i)[0]) and ":" in str(file_data.row_values(i)[0]):
                                norm.append(file_data.row_values(i)[0].split(":")[1].strip())  # strip()的作用是去除前后空格
                                norm_row.append(i)

                    nameindex = 0
                    concindex = 0
                    # 第一个化合物表格确定samplename和浓度所在列，norm_row[0]为第一个化合物所在行，+2是该化合物表格位于该化合物所在行的下两行
                    for i in range(len(file_data.row_values(norm_row[0]+2))):
                        if file_data.row_values(norm_row[0]+2)[i] == "ID":
                            nameindex = i
                        elif "nmol/L" in file_data.row_values(norm_row[0]+2)[i]:
                            concindex = i

                    # 未规范设置表格列名
                    if nameindex == 0 or concindex == 0:
                        error_message = "未规范设置文件列名，请检查并规范设置后重新提交数据!"
                        return {"error_message": error_message}

                    # 确定本底数(含有"RecB")
                    RecBlist = []  # 本底列表,长度/3即为本底数
                    if len(norm) == 1:  # 如果只有一个化合物,则循环第一个化合物所在行到最后一行
                        for i in range(norm_row[0], nrows):
                            if "RecB" in file_data.row_values(i)[nameindex][0:6] and file_data.row_values(i)[nameindex][0:6] not in RecBlist:
                                RecBlist.append(file_data.row_values(i)[nameindex][0:6])
                    else:  # 如果有多个化合物,则循环第一个化合物所在行到第二个化合物所在行
                        for i in range(norm_row[0], norm_row[1]):
                            if "RecB" in file_data.row_values(i)[nameindex][0:6] and file_data.row_values(i)[nameindex][0:6] not in RecBlist:
                                RecBlist.append(file_data.row_values(i)[nameindex][0:6])

                    for k in range(len(norm)):

                        Recycle_enddict_show[norm[k]] = {}  # 数据展示字典
                        Recycle_enddict_savedata[norm[k]] = {}  # 数据保存字典

                        for j in range(len(RecBlist)):  # 本底列表
                            RecB_conc = []  # 本底浓度列表
                            low = []  # 本底加标后低浓度列表
                            median = []  # 本底加标后中浓度列表
                            high = []  # 本底加标后高浓度列表
                            if k < len(norm)-1:  # 如果不是最后一个化合物，索引为该化合物所在行到后一个化合物所在行
                                for i in range(norm_row[k], norm_row[k+1]):
                                    # 如果实验号命名方式匹配上，则在相应列表中添加相应数据
                                    if RecBlist[j] in file_data.row_values(i)[nameindex]:
                                        RecB_conc.append(effectnum(file_data.row_values(i)[concindex], digits))
                                    elif "L" in file_data.row_values(i)[nameindex] and Reclist[j] in file_data.row_values(i)[nameindex]:
                                        low.append(effectnum(file_data.row_values(i)[concindex], digits))
                                    elif "M" in file_data.row_values(i)[nameindex] and Reclist[j] in file_data.row_values(i)[nameindex]:
                                        median.append(effectnum(file_data.row_values(i)[concindex], digits))
                                    elif "H" in file_data.row_values(i)[nameindex] and Reclist[j] in file_data.row_values(i)[nameindex]:
                                        high.append(effectnum(file_data.row_values(i)[concindex], digits))
                            else:
                                for i in range(norm_row[k], nrows):
                                    # 如果实验号命名方式匹配上，则在相应列表中添加相应数据
                                    if RecBlist[j] in file_data.row_values(i)[nameindex]:
                                        RecB_conc.append(effectnum(file_data.row_values(i)[concindex], digits))
                                    elif "L" in file_data.row_values(i)[nameindex] and Reclist[j] in file_data.row_values(i)[nameindex]:
                                        low.append(effectnum(file_data.row_values(i)[concindex], digits))
                                    elif "M" in file_data.row_values(i)[nameindex] and Reclist[j] in file_data.row_values(i)[nameindex]:
                                        median.append(effectnum(file_data.row_values(i)[concindex], digits))
                                    elif "H" in file_data.row_values(i)[nameindex] and Reclist[j] in file_data.row_values(i)[nameindex]:
                                        high.append(effectnum(file_data.row_values(i)[concindex], digits))

                            # 数据展示字典
                            Recycle_enddict_show[norm[k]][RecBlist[j]] = []
                            Recycle_enddict_show[norm[k]][RecBlist[j]].extend(RecB_conc)
                            Recycle_enddict_show[norm[k]][RecBlist[j]].extend(low)
                            Recycle_enddict_show[norm[k]][RecBlist[j]].extend(median)
                            Recycle_enddict_show[norm[k]][RecBlist[j]].extend(high)

                            if Rec_dict != {}:
                                Recycle_enddict_show[norm[k]][RecBlist[j]].extend(Rec_dict[Reclist[j]])
                            else:
                                Recycle_enddict_show[norm[k]][RecBlist[j]].extend(["", "", "", "", "", "", "", "", ""])

                            # 数据保存字典
                            Recycle_enddict_savedata[norm[k]][RecBlist[j]] = []
                            Recycle_enddict_savedata[norm[k]][RecBlist[j]].extend(RecB_conc)
                            Recycle_enddict_savedata[norm[k]][RecBlist[j]].extend(low)
                            Recycle_enddict_savedata[norm[k]][RecBlist[j]].extend(median)
                            Recycle_enddict_savedata[norm[k]][RecBlist[j]].extend(high)

                elif manufacturers == "Thermo":
                    data = xlrd.open_workbook(filename=None, file_contents=file.read())  # 读取表格
                    norm = []  # Thermo的原始数据格式为一个化合物一个sheet,获取每个sheet的名字,即为化合物名称
                    sheetindex = []  # 需要的化合物所在sheet索引列表
                    for index in range(len(data.sheet_names())):
                        norm.append(data.sheet_names()[index].split("Sheet1")[1])
                        sheetindex.append(index)

                    # 循环读取每个sheet工作表,即为每个化合物的表
                    for index in range(len(sheetindex)):
                        file_data = data.sheets()[sheetindex[index]]
                        nrows = file_data.nrows
                        ncols = file_data.ncols

                        # 第一行确定samplename和浓度所在列
                        nameindex = 0
                        concindex = 0
                        for i in range(len(file_data.row_values(8))):
                            if file_data.row_values(8)[i] == "Raw File Name":
                                nameindex = i
                            elif file_data.row_values(8)[i] == "Calculated Amount":
                                concindex = i

                        # 未规范设置表格列名
                        if concindex == 0:
                            error_message = "未规范设置文件列名，请检查并规范设置后重新提交数据!"
                            return {"error_message": error_message}

                        # 确定本底数(含有"RecB")
                        RecBlist = []  # 本底列表,长度/3即为本底数
                        for i in range(nrows):
                            if "RecB" in file_data.row_values(i)[nameindex][0:6] and file_data.row_values(i)[nameindex][0:6] not in RecBlist:
                                RecBlist.append(file_data.row_values(i)[nameindex][0:6])
                        
                        # 匹配原始数据中与加标回收相关的行
                        Recycle_enddict_show[norm[index]] = {}  # 数据展示字典
                        Recycle_enddict_savedata[norm[index]] = {}  # 数据保存字典

                        for j in range(len(RecBlist)):  # 本底列表
                            RecB_conc = []  # 本底浓度列表
                            low = []  # 本底加标后低浓度列表
                            median = []  # 本底加标后中浓度列表
                            high = []  # 本底加标后高浓度列表
                            for i in range(nrows):
                                # 如果实验号命名方式匹配上，则在相应列表中添加相应数据
                                if RecBlist[j] in file_data.row_values(i)[nameindex]:
                                    RecB_conc.append(effectnum(file_data.row_values(i)[concindex], digits))
                                elif "L" in file_data.row_values(i)[nameindex] and Reclist[j] in file_data.row_values(i)[nameindex]:
                                    low.append(effectnum(file_data.row_values(i)[concindex], digits))
                                elif "M" in file_data.row_values(i)[nameindex] and Reclist[j] in file_data.row_values(i)[nameindex]:
                                    median.append(effectnum(file_data.row_values(i)[concindex], digits))
                                elif "H" in file_data.row_values(i)[nameindex] and Reclist[j] in file_data.row_values(i)[nameindex]:
                                    high.append(effectnum(file_data.row_values(i)[concindex], digits))

                            # 数据展示字典
                            Recycle_enddict_show[norm[index]][RecBlist[j]] = []
                            Recycle_enddict_show[norm[index]][RecBlist[j]].extend(RecB_conc)
                            Recycle_enddict_show[norm[index]][RecBlist[j]].extend(low)
                            Recycle_enddict_show[norm[index]][RecBlist[j]].extend(median)
                            Recycle_enddict_show[norm[index]][RecBlist[j]].extend(high)

                            if Rec_dict != {}:
                                Recycle_enddict_show[norm[index]][RecBlist[j]].extend(Rec_dict[Reclist[j]])
                            else:
                                Recycle_enddict_show[norm[index]][RecBlist[j]].extend(["", "", "", "", "", "", "", "", ""])

                            # 数据保存字典
                            Recycle_enddict_savedata[norm[index]][RecBlist[j]] = []
                            Recycle_enddict_savedata[norm[index]][RecBlist[j]].extend(RecB_conc)
                            Recycle_enddict_savedata[norm[index]][RecBlist[j]].extend(low)
                            Recycle_enddict_savedata[norm[index]][RecBlist[j]].extend(median)
                            Recycle_enddict_savedata[norm[index]][RecBlist[j]].extend(high)

                elif manufacturers == "岛津":
                    # 读取txt
                    content = []
                    for line in file:
                        content.append(line.decode("GB2312").replace("\r\n", "").split("\t"))

                    nameindex = 0
                    concindexx = 0  # 浓度索引，岛津的数据格式决定每个化合物的浓度所在列一定是同一列
                    norm = []  # 化合物列表
                    norm_row = []  # 化合物所在行

                    for i in range(len(content)):
                        if content[i][0] == "Name":  # 如果某一行第一列为"Name"，则该行第二列为化合物名称

                            # 若化合物名称后含有“-”，需切除
                            if "-" in content[i][1]:
                                norm.append(content[i][1].split("-")[0])
                            else:
                                norm.append(content[i][1])
                            norm_row.append(i)

                    for i in range(len(content[2])):  # 第二行确定samplename和浓度所在列
                        if content[2][i] == "数据文件名":
                            nameindex = i
                        elif content[2][i] == "浓度":
                            concindex = i

                    # 未规范设置表格列名
                    if nameindex == 0 or concindex == 0:
                        error_message = "未规范设置文件列名，请检查并规范设置后重新提交数据!"
                        return {"error_message": error_message}

                    # 确定本底数(含有"RecB")
                    RecBlist = []  # 本底列表,长度/3即为本底数
                    if len(norm) == 1:  # 如果只有一个化合物,则循环第一个化合物所在行到最后一行
                        for i in range(norm_row[0], len(content)-1):
                            if "RecB" in file_data.row_values(i)[nameindex][0:6] and file_data.row_values(i)[nameindex][0:6] not in RecBlist:
                                RecBlist.append(file_data.row_values(i)[nameindex][0:6])
                    else:  # 如果有多个化合物,则循环第一个化合物所在行到第二个化合物所在行
                        for i in range(norm_row[0], norm_row[1]):
                            if "RecB" in file_data.row_values(i)[nameindex][0:6] and file_data.row_values(i)[nameindex][0:6] not in RecBlist:
                                RecBlist.append(file_data.row_values(i)[nameindex][0:6])

                    # 匹配原始数据中与加标回收相关的行
                    for k in range(len(norm)):

                        Recycle_enddict_show[norm[k]] = {}  # 数据展示字典
                        Recycle_enddict_savedata[norm[k]] = {}  # 数据保存字典

                        for j in range(len(RecBlist)):  # 本底列表
                            RecB_conc = []  # 本底浓度列表
                            low = []  # 本底加标后低浓度列表
                            median = []  # 本底加标后中浓度列表
                            high = []  # 本底加标后高浓度列表
                            if k < len(norm)-1:  # 如果不是最后一个化合物，索引为该化合物所在行到后一个化合物所在行
                                for i in range(norm_row[k], norm_row[k+1]):
                                    # 如果实验号命名方式匹配上，则在相应列表中添加相应数据
                                    if RecBlist[j] in content[i][nameindex]:
                                        RecB_conc.append(effectnum(content[i][concindex], digits))
                                    elif "L" in content[i][nameindex] and Reclist[j] in content[i][nameindex]:
                                        low.append(effectnum(content[i][concindex], digits))
                                    elif "M" in content[i][nameindex] and Reclist[j] in content[i][nameindex]:
                                        median.append(effectnum(content[i][concindex], digits))
                                    elif "H" in content[i][nameindex] and Reclist[j] in content[i][nameindex]:
                                        high.append(effectnum(content[i][concindex], digits))

                            else:
                                for i in range(norm_row[k], len(content)-1):
                                    # 如果实验号命名方式匹配上，则在相应列表中添加相应数据
                                    if RecBlist[j] in content[i][nameindex]:
                                        RecB_conc.append(effectnum(content[i][concindex], digits))
                                    elif "L" in content[i][nameindex] and Reclist[j] in content[i][nameindex]:
                                        low.append(effectnum(content[i][concindex], digits))
                                    elif "M" in content[i][nameindex] and Reclist[j] in content[i][nameindex]:
                                        median.append(effectnum(content[i][concindex], digits))
                                    elif "H" in content[i][nameindex] and Reclist[j] in content[i][nameindex]:
                                        high.append(effectnum(content[i][concindex], digits))

                            # 数据展示字典
                            Recycle_enddict_show[norm[k]][RecBlist[j]] = []
                            Recycle_enddict_show[norm[k]][RecBlist[j]].extend(RecB_conc)
                            Recycle_enddict_show[norm[k]][RecBlist[j]].extend(low)
                            Recycle_enddict_show[norm[k]][RecBlist[j]].extend(median)
                            Recycle_enddict_show[norm[k]][RecBlist[j]].extend(high)

                            if Rec_dict != {}:
                                Recycle_enddict_show[norm[k]][RecBlist[j]].extend(Rec_dict[Reclist[j]])
                            else:
                                Recycle_enddict_show[norm[k]][RecBlist[j]].extend(["", "", "", "", "", "", "", "", ""])

                            # 数据保存字典
                            Recycle_enddict_savedata[norm[k]][RecBlist[j]] = []
                            Recycle_enddict_savedata[norm[k]][RecBlist[j]].extend(RecB_conc)
                            Recycle_enddict_savedata[norm[k]][RecBlist[j]].extend(low)
                            Recycle_enddict_savedata[norm[k]][RecBlist[j]].extend(median)
                            Recycle_enddict_savedata[norm[k]][RecBlist[j]].extend(high)

                elif manufacturers == "AB":
                    # 定义化合物列表，列表统一命名为norm
                    norm = normAB

                    # 获取上传的文件
                    file_data = Document(file)

                    # 每个表格最上方的标题内容列表，含有母离子和子离子的信息。需依此及母离子和子离子列表判断table索引
                    paragraphs = []

                    # 若标题长度为0或为换行等，不添加进入标题内容列表
                    for p in file_data.paragraphs:
                        if len(p.text) != 0 and p.text != "\n" and len(p.text.strip()) != 0:
                            paragraphs.append(p.text.strip())

                    # 确定table索引，母离子和子离子都与后台管理系统中设置的数值相同才证明是需要读取的定量表格
                    tableindex = []
                    for i in range(len(paragraphs)):
                        for j in range(len(ZP_Method_precursor_ion)):
                            if ZP_Method_precursor_ion[j] in paragraphs[i] and ZP_Method_product_ion[j] in paragraphs[i]:
                                tableindex.append(2*i+1)

                    tables = file_data.tables  # 获取文件中的表格集

                    # 循环定量表格的索引
                    for k in range(len(tableindex)):
                        # 获取文件中的定量表格
                        tablequantify = tables[tableindex[k]]

                        # 先把表格里的所有数据取出来放进一个列表中，读取速度会比直接读表格快很多
                        cells = tablequantify._cells
                        ROWS = len(tablequantify.rows)
                        COLUMNS = len(tablequantify.columns)
                        rowdatalist = []  # 每一行的数据列表
                        rowdatagatherlist = []  # 每一行的数据列表汇总列表

                        for i in range(ROWS*COLUMNS):
                            text = cells[i].text.replace("\n", "")
                            text = text.strip()  # 去除空白符
                            if i % COLUMNS != 0 or i == 0:
                                rowdatalist.append(text)
                            else:
                                rowdatagatherlist.append(rowdatalist)
                                rowdatalist = []
                                rowdatalist.append(text)
                        rowdatagatherlist.append(rowdatalist)

                        nameindex = 0
                        concindex = 0  # 浓度索引，AB的数据格式决定每个化合物的浓度所在列一定是同一列

                        # 读取表格的第一行的单元格,判断实验号和浓度索引
                        for i in range(len(rowdatagatherlist[0])):
                            if rowdatagatherlist[0][i] == "Sample Name":
                                nameindex = i
                            elif "Calculated Conc" in rowdatagatherlist[0][i]:
                                concindex = i

                        # 确定本底数(含有"RecB")
                        RecBlist = []  # 本底列表,长度/3即为本底数
                        for i in range(len(rowdatagatherlist)):  # 循环原始数据中的每一行,并避免重复添加
                            # 加标回收率本底实验号命名“RecB-1-1”
                            if "RecB" in rowdatagatherlist[i][nameindex][0:6] and rowdatagatherlist[i][nameindex][0:6] not in RecBlist:
                                RecBlist.append(rowdatagatherlist[i][nameindex][0:6])

                        Recycle_enddict_show[norm[k]] = {}  # 数据展示字典
                        Recycle_enddict_savedata[norm[k]] = {}  # 数据保存字典

                        for j in range(len(RecBlist)):  # 本底列表
                            RecB_conc = []  # 本底浓度列表
                            low = []  # 本底加标后低浓度列表
                            median = []  # 本底加标后中浓度列表
                            high = []  # 本底加标后高浓度列表
                            for i in range(len(rowdatagatherlist)):  # 循环原始数据中的每一行
                                if RecBlist[j] in rowdatagatherlist[i][nameindex]:
                                    # 判断浓度是否为N/A,本底浓度为N/A，添加0
                                    if rowdatagatherlist[i][concindex]=="N/A":
                                        RecB_conc.append("0")
                                    else:
                                        RecB_conc.append(effectnum(rowdatagatherlist[i][concindex], digits))
                                elif "L" in rowdatagatherlist[i][nameindex] and Reclist[j] in rowdatagatherlist[i][nameindex]:
                                    # 判断浓度是否为N/A
                                    if rowdatagatherlist[i][concindex]=="N/A":
                                        low.append("N/A")
                                    else:
                                        low.append(effectnum(rowdatagatherlist[i][concindex], digits))
                                elif "M" in rowdatagatherlist[i][nameindex] and Reclist[j] in rowdatagatherlist[i][nameindex]:
                                    # 判断浓度是否为N/A
                                    if rowdatagatherlist[i][concindex]=="N/A":
                                        median.append("N/A")
                                    else:
                                        median.append(effectnum(rowdatagatherlist[i][concindex], digits))
                                elif "H" in rowdatagatherlist[i][nameindex] and Reclist[j] in rowdatagatherlist[i][nameindex]:
                                    # 判断浓度是否为N/A
                                    if rowdatagatherlist[i][concindex]=="N/A":
                                        high.append("N/A")
                                    else:
                                        high.append(effectnum(rowdatagatherlist[i][concindex], digits))

                            # 数据展示字典
                            Recycle_enddict_show[norm[k]][RecBlist[j]] = []
                            Recycle_enddict_show[norm[k]][RecBlist[j]].extend(RecB_conc)
                            Recycle_enddict_show[norm[k]][RecBlist[j]].extend(low)
                            Recycle_enddict_show[norm[k]][RecBlist[j]].extend(median)
                            Recycle_enddict_show[norm[k]][RecBlist[j]].extend(high)

                            if Rec_dict != {}:
                                Recycle_enddict_show[norm[k]][RecBlist[j]].extend(Rec_dict[Reclist[j]])
                            else:
                                Recycle_enddict_show[norm[k]][RecBlist[j]].extend(["", "", "", "", "", "", "", "", ""])

                            # 数据保存字典
                            Recycle_enddict_savedata[norm[k]][RecBlist[j]] = []
                            Recycle_enddict_savedata[norm[k]][RecBlist[j]].extend(RecB_conc)
                            Recycle_enddict_savedata[norm[k]][RecBlist[j]].extend(low)
                            Recycle_enddict_savedata[norm[k]][RecBlist[j]].extend(median)
                            Recycle_enddict_savedata[norm[k]][RecBlist[j]].extend(high)

            elif platform == "液相":
                if manufacturers == "Agilent":
                    data = xlrd.open_workbook(filename=None, file_contents=file.read())  # 读取表格
                    file_data = data.sheets()[0]
                    nrows = file_data.nrows
                    ncols = file_data.ncols

                    norm = []  # 化合物列表
                    norm_row = []  # 化合物所在行
                    for j in range(nrows):
                        # 如果某一行的第一个元素为“化合物”，则添加第三个元素进入化合物列表
                        if file_data.row_values(j)[0] == "化合物:":
                            norm.append(file_data.row_values(j)[2])
                            norm_row.append(j)

                    nameindex = 0
                    conindex = 0
                    # 第一个化合物表格确定samplename和浓度所在列，norm_row[0]为第一个化合物所在行，+2是该化合物表格位于该化合物所在行的下两行
                    for i in range(len(file_data.row_values(norm_row[0]+2))):
                        if file_data.row_values(norm_row[0]+2)[i] == "样品名称":
                            nameindex = i
                        elif "含量" in file_data.row_values(norm_row[0]+2)[i]:
                            conindex = i

                    # 确定本底数，含有"Recycle"及"background"(以第一个化合物为准确定本底数)
                    background = []  # 本底列表,长度/3即为本底数
                    if len(norm) == 1:  # 如果只有一个化合物,则循环第一个化合物所在行到最后一行
                        for i in range(norm_row[0], nrows):
                            if "Recycle" in file_data.row_values(i)[nameindex] and "background" in file_data.row_values(i)[nameindex]:
                                background.append(
                                    file_data.row_values(i)[nameindex])

                    else:  # 如果有多个化合物,则循环第一个化合物所在行到第二个化合物所在行
                        for i in range(norm_row[0], norm_row[1]):
                            if "Recycle" in file_data.row_values(i)[nameindex] and "background" in file_data.row_values(i)[nameindex]:
                                background.append(
                                    file_data.row_values(i)[nameindex])

                    # 匹配原始数据中与加标回收相关的行
                    for k in range(len(norm)):
                        group_Recycle = {}  # 每个化合物数据字典
                        for j in range(int(len(background)/3)):  # 本底列表,长度/3即为本底数
                            background_conc = []  # 本底浓度列表
                            low = []  # 本底加标后低浓度列表
                            median = []  # 本底加标后中浓度列表
                            high = []  # 本底加标后高浓度列表
                            if k < len(norm)-1:  # 如果不是最后一个化合物，索引为该化合物所在行到后一个化合物所在行
                                for i in range(norm_row[k], norm_row[k+1]):
                                    # 如果实验号命名方式匹配上，则在相应列表中添加相应数据
                                    if "background" in file_data.row_values(i)[nameindex] and Recycle_background[j] in file_data.row_values(i)[nameindex]:
                                        background_conc.append(
                                            effectnum(file_data.row_values(i)[conindex], digits))
                                    elif "low" in file_data.row_values(i)[nameindex] and Recycle_background[j] in file_data.row_values(i)[nameindex]:
                                        low.append(
                                            effectnum(file_data.row_values(i)[conindex], digits))
                                    elif "median" in file_data.row_values(i)[nameindex] and Recycle_background[j] in file_data.row_values(i)[nameindex]:
                                        median.append(
                                            effectnum(file_data.row_values(i)[conindex], digits))
                                    elif "high" in file_data.row_values(i)[nameindex] and Recycle_background[j] in file_data.row_values(i)[nameindex]:
                                        high.append(
                                            effectnum(file_data.row_values(i)[conindex], digits))
                            else:
                                for i in range(norm_row[k], nrows):
                                    # 如果实验号命名方式匹配上，则在相应列表中添加相应数据
                                    if "background" in file_data.row_values(i)[nameindex] and Recycle_background[j] in file_data.row_values(i)[nameindex]:
                                        background_conc.append(
                                            effectnum(file_data.row_values(i)[conindex], digits))
                                    elif "low" in file_data.row_values(i)[nameindex] and Recycle_background[j] in file_data.row_values(i)[nameindex]:
                                        low.append(
                                            effectnum(file_data.row_values(i)[conindex], digits))
                                    elif "median" in file_data.row_values(i)[nameindex] and Recycle_background[j] in file_data.row_values(i)[nameindex]:
                                        median.append(
                                            effectnum(file_data.row_values(i)[conindex], digits))
                                    elif "high" in file_data.row_values(i)[nameindex] and Recycle_background[j] in file_data.row_values(i)[nameindex]:
                                        high.append(
                                            effectnum(file_data.row_values(i)[conindex], digits))

                            group_Recycle[Recycle_background[j]] = []
                            for i in background_conc:
                                group_Recycle[Recycle_background[j]].append(i)
                            for i in low:
                                group_Recycle[Recycle_background[j]].append(i)
                            for i in median:
                                group_Recycle[Recycle_background[j]].append(i)
                            for i in high:
                                group_Recycle[Recycle_background[j]].append(i)
                        Recycle_enddict_show[norm[k]] = group_Recycle

                    print(Recycle_enddict)

    # 删除Recycle_enddict_show和Recycle_enddict_savedata未做加标回收的化合物
    print(Recycle_enddict_show)
    print(Recycle_enddict_savedata)


    return {"Recycle_enddict_show": Recycle_enddict_show, "Recycle_enddict_savedata": Recycle_enddict_savedata, "Unit": Unit, "lowvalue": lowvalue, "upvalue": upvalue}


# PT数据关联进入最终报告
def related_PT(id):
    # 第一步：后台描述性内容数据提取
    # 1 根据id找到项目
    project = ReportInfo.objects.get(id=id).project

    # 2 优先查找特殊参数设置里是否有数据，如有就调用，没有则调用通用性参数设置里的数据
    # 特殊参数设置描述性内容
    textlist_special = []
    try:
        special_1 = Special.objects.get(project=project)
        special_2 = PTspecial.objects.get(special=special_1)
        if PTspecialtexts.objects.filter(pTspecial=special_2).count() > 0:
            text_special = PTspecialtexts.objects.filter(pTspecial=special_2)
            for i in text_special:
                textlist_special.append(i.text)
    except:
        pass

    # 3 通用数据抓取
    # 描述性内容
    textlist_general = []
    general_1 = General.objects.get(name="通用性项目")  # 通用参数设置描述性内容
    general_2 = PTgeneral.objects.get(general=general_1)
    text_general = PTgeneraltexts.objects.filter(pTgeneral=general_2)
    for i in text_general:
        textlist_general.append(i.text)

    # 查找是否单独设置了每个化合物的有效位数
    DIGITS_TABLE = Special.objects.get(project=project)
    pt_special = PTspecial.objects.get(special=DIGITS_TABLE)
    pt_accept = PTspecialaccept.objects.filter(pTspecial=pt_special)
    Digitslist = []  # 每个化合物有效位数列表
    Digitsdict = {}  # 每个化合物有效位数字典

    for i in pt_accept:
        Digitslist.append(i.digits)

    if Digitslist == [] or Digitslist[0] == "":  # 如果全部没设置或者只是单位没设置
        pass
    else:
        for i in pt_accept:
            Digitsdict[i.norm] = i.digits

    # 第二步：报告数据提取

    # 定义需要生成的字典
    PT_dict = {}  # 最终需要的字典
    targetcolname = ""

    # 1 基础数据抓取
    PT_data = PT.objects.filter(reportinfo_id=id)

    PT_norm = []
    for i in PT_data:
        if i.norm not in PT_norm:
            PT_norm.append(i.norm)

    # 判断是模板1(可接受区间)还是模板2(可接受标准)
    PT_templates = []
    for i in PT_data:
        if int(i.templates) not in PT_templates:
            PT_templates.append(int(i.templates))

    for i in PT_norm:
        middle_list = []  # 每个化合物的数据列表
        middle_table = PT.objects.filter(reportinfo_id=id, norm=i)

        # 可接受区间模板
        if PT_templates[0] == 1:

            # 1 非25OHD项目
            if "25OHD" not in project:
                for j in middle_table:
                    # 判断是否有不需要显示的数据
                    if j.accept1=="不显示":
                        pass
                    else:
                        # # 没有为每个化合物单独设置有效位数，则调用通用性设置
                        # if Digitsdict == {} or list(Digitsdict.values())[0] == None:
                        #     rowlist = []  # 每一行的小列表
                        #     rowlist.append(j.Experimentnum)
                        #     rowlist.append(j.value)
                        #     rowlist.append(j.accept1)
                        #     rowlist.append(j.accept2)
                        #     rowlist.append(j.PT_pass)
                        #     middle_list.append(rowlist)
                        # # 为每个化合物单独设置了有效位数，则调用每个化合物的设置
                        # else:
                        #     rowlist = []
                        #     rowlist.append(j.Experimentnum)
                        #     rowlist.append(effectnum(j.value, Digitsdict[i]))
                        #     rowlist.append(j.accept1)
                        #     rowlist.append(j.accept2)
                        #     rowlist.append(j.PT_pass)
                        #     middle_list.append(rowlist)

                        try:
                            rowlist = []
                            rowlist.append(j.Experimentnum)
                            rowlist.append(effectnum(j.value, Digitsdict[i]))
                            rowlist.append(effectnum(j.accept1, Digitsdict[i]))
                            rowlist.append(effectnum(j.accept2, Digitsdict[i]))
                            rowlist.append(j.PT_pass)
                            middle_list.append(rowlist)
                        except:
                            rowlist = []  # 每一行的小列表
                            rowlist.append(j.Experimentnum)
                            rowlist.append(j.value)
                            rowlist.append(j.accept1)
                            rowlist.append(j.accept2)
                            rowlist.append(j.PT_pass)
                            middle_list.append(rowlist)
                PT_dict[i] = middle_list

            else:
                pass
        
        # 可接受标准模板
        else:
            # 1 非25OHD项目
            if "25OHD" not in project:
                for j in middle_table:
                    # 判断是否有不需要显示的数据
                    if j.target=="不显示":
                        pass
                    else:
                        # 没有为每个化合物单独设置有效位数，则调用通用性设置
                        if Digitsdict == {} or list(Digitsdict.values())[0] == None:
                            rowlist = []  # 每一行的小列表
                            rowlist.append(j.Experimentnum)
                            rowlist.append(j.value)

                            if "-" in j.target:
                                middlevalue = j.target
                                rowlist.append(middlevalue.split("-")[0])
                                targetcolname = middlevalue.split("-")[1]
                            else:
                                rowlist.append(j.target)
                            rowlist.append(j.received)
                            rowlist.append(j.bias)
                            rowlist.append(j.PT_pass)
                            middle_list.append(rowlist)
                        # 为每个化合物单独设置了有效位数，则调用每个化合物的设置
                        else:
                            rowlist = []
                            rowlist.append(j.Experimentnum)
                            rowlist.append(effectnum(j.value, Digitsdict[i]))
                            if "-" in j.target:
                                middlevalue = j.target
                                rowlist.append(middlevalue.split("-")[0])
                                targetcolname = middlevalue.split("-")[1]
                            else:
                                rowlist.append(j.target)
                            rowlist.append(j.received)
                            rowlist.append(j.bias)
                            rowlist.append(j.PT_pass)
                            middle_list.append(rowlist)
                PT_dict[i] = middle_list

            else:
                for j in middle_table:
                    # 判断是否有不需要显示的数据
                    if j.target=="不显示":
                        pass
                    else:
                        # 没有为每个化合物单独设置有效位数，则调用通用性设置
                        if Digitsdict == {} or list(Digitsdict.values())[0] == None:
                            rowlist = []  # 每一行的小列表
                            rowlist.append(j.Experimentnum)
                            rowlist.append(j.value.split("-")[0])  # 添加D2结果
                            rowlist.append(j.value.split("-")[1])  # 添加D3结果
                            rowlist.append(j.value.split("-")[2])  # 添加总D结果
                            rowlist.append(j.target)
                            rowlist.append(j.received)
                            rowlist.append(j.bias)
                            rowlist.append(j.PT_pass)
                            middle_list.append(rowlist)
                        # 为每个化合物单独设置了有效位数，则调用每个化合物的设置
                        else:
                            rowlist = []
                            rowlist.append(j.Experimentnum)
                            rowlist.append(effectnum(j.value, Digitsdict[i]))
                            rowlist.append(j.target)
                            rowlist.append(j.received)
                            rowlist.append(j.bias)
                            rowlist.append(j.PT_pass)
                            middle_list.append(rowlist)
                PT_dict[i] = middle_list

    # 将PT_dict中列表为空的元素删除,并记录删除的个数，以便修改前端表格的索引
    delcount=0
    for key in list(PT_dict.keys()):
        if PT_dict[key]==[]:
            del PT_dict[key]
            delcount+=1
    
    print(PT_dict)
    print(targetcolname)
    if len(textlist_special) != 0:
        return {"PT_dict": PT_dict, "PT_templates": PT_templates, "textlist": textlist_special, "serial": len(textlist_special)+1,"delcount":delcount,"targetcolname":targetcolname}
    else:
        return {"PT_dict": PT_dict, "PT_templates": PT_templates, "textlist": textlist_general, "serial": len(textlist_general)+1,"delcount":delcount,"targetcolname":targetcolname}


# 加标回收据关联进入最终报告
def related_recycle(id):
    # 第一步：后台描述性内容数据提取

    # 根据id找到项目
    project = ReportInfo.objects.get(id=id).project

    # 优先查找特殊参数设置里是否有数据，如有就调用，没有则调用通用性参数设置里的数据
    # 特殊数据抓取
    zqd_special = Special.objects.get(project=project)
    Recycle_special = Recyclespecial.objects.get(special=zqd_special)
    textlist_special = []  # 特殊参数设置描述性内容
    if Recyclespecialtexts.objects.filter(recyclespecial=Recycle_special).count() > 0:
        text_special = Recyclespecialtexts.objects.filter(recyclespecial=Recycle_special)
        for i in text_special:
            textlist_special.append(i.text)

    # 通用数据抓取
    zqd_general = General.objects.get(name="通用性项目")
    Recycle_general = Recyclegeneral.objects.get(general=zqd_general)
    text_general = Recyclegeneraltexts.objects.filter(recyclegeneral=Recycle_general)
    textlist_general = []
    for i in text_general:
        textlist_general.append(i.text)

    # 查找是否单独设置了每个化合物的有效位数
    DIGITS_TABLE = Special.objects.get(project=project)
    pt_special = PTspecial.objects.get(special=DIGITS_TABLE)
    pt_accept = PTspecialaccept.objects.filter(pTspecial=pt_special)
    Digitslist = []  # 每个化合物有效位数列表
    Digitsdict = {}  # 每个化合物有效位数字典

    for i in pt_accept:
        Digitslist.append(i.digits)

    if Digitslist == [] or Digitslist[0] == "":  # 如果全部没设置或者只是单位没设置
        pass
    else:
        for i in pt_accept:
            Digitsdict[i.norm] = i.digits

    # 第二步：报告数据提取

    # 定义需要生成的字典
    Recycle_dict = {}  # 最终需要的字典
    try:
        # 1 基础数据抓取
        Recycle_data = RECYCLE.objects.filter(reportinfo_id=id)

        Recycle_norm = []  # 去重后的指标列表
        Recycle_sam = []  # 去重后的本底列表

        for item in Recycle_data:
            if item.norm not in Recycle_norm:
                Recycle_norm.append(item.norm)
            if item.Experimentnum not in Recycle_sam:
                Recycle_sam.append(item.Experimentnum)

        # Recycle_endconclusion = "" # 最终结论
        for i in Recycle_norm:
            middle_dict = {}  # 每个化合物的字典
            for j in Recycle_sam:
                middle_table = RECYCLE.objects.filter(reportinfo_id=id, norm=i, Experimentnum=j)
                sum = 0 # 计算平均检测浓度用
                middle_sam = []
                # 每个化合物的回收率列表，方便提取最大最小值  "maxrecycle":str(max(recycle_list))+"%","minrecycle":str(min(recycle_list))+"%"
                middle_list = []
                for item in middle_table:
                    # 判断是否有不需要显示的数据
                    print(item.theory_conc)
                    if item.theory_conc=="不显示":
                        pass
                    
                    else:
                        # 没有为每个化合物单独设置有效位数，则调用通用性设置
                        if Digitsdict == {} or list(Digitsdict.values())[0] == None:
                            middle_sam.append(item.sam_conc)
                            sum += float(item.sam_conc)
                            middle_sam.append(item.theory_conc)
                            middle_sam.append(item.level)
                            middle_sam.append(item.end_conc1)
                            middle_sam.append(item.end_conc2)
                            middle_sam.append(item.end_conc3)
                            middle_sam.append(item.end_recycle1)
                            middle_sam.append(item.end_recycle2)
                            middle_sam.append(item.end_recycle3)
                            middle_list.append(float(item.end_recycle1))
                            middle_list.append(float(item.end_recycle2))
                            middle_list.append(float(item.end_recycle3))
                        # 为每个化合物单独设置了有效位数，则调用每个化合物的设置
                        else:
                            middle_sam.append(effectnum(item.sam_conc, Digitsdict[i]))
                            sum += float(item.sam_conc)
                            middle_sam.append(item.theory_conc)
                            middle_sam.append(item.level)
                            middle_sam.append(effectnum(item.end_conc1, Digitsdict[i]))
                            middle_sam.append(effectnum(item.end_conc2, Digitsdict[i]))
                            middle_sam.append(effectnum(item.end_conc3, Digitsdict[i]))
                            middle_sam.append(item.end_recycle1)
                            middle_sam.append(item.end_recycle2)
                            middle_sam.append(item.end_recycle3)
                            middle_list.append(float(item.end_recycle1))
                            middle_list.append(float(item.end_recycle2))
                            middle_list.append(float(item.end_recycle3))

                if middle_sam!=[]:
                    middle_sam.append(new_round(sum/3, 1))
                    middle_dict[j] = middle_sam
            Recycle_dict[i] = middle_dict

        # 将Recycle_dict中列表为空的元素删除,并记录删除的个数，以便修改前端表格的索引
        delcount=0
        for key in list(Recycle_dict.keys()):
            if Recycle_dict[key]=={}:
                del Recycle_dict[key]
                delcount+=1

        # print(Recycle_dict)
        if len(textlist_special) != 0:
            return {"Recycle_dict": Recycle_dict, "textlist": textlist_special, "serial": len(textlist_special)+1,"delcount":delcount}
        else:
            return {"Recycle_dict": Recycle_dict, "textlist": textlist_general, "serial": len(textlist_general)+1,"delcount":delcount}


    except:
        pass