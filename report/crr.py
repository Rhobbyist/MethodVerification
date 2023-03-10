from django.http import HttpResponse
from django.http import HttpResponseRedirect
from report import models
from report.models import *
import xlrd
import numpy as np
import math
from report.effectnum import *
from docx import Document
import re


def CRRfileread(files, reportinfo, project, platform, manufacturers,Detectionplatform, Unit, digits, ZP_Method_precursor_ion, ZP_Method_product_ion, normAB, Number_of_compounds):

    # 第一步:后台数据抓取（回收率上下限）
    id1 = Special.objects.get(project=project).id
    id2 = CRRspecial.objects.get(special_id=id1).id

    # 优先查找特殊参数设置里是否有数据，如有就调用，没有则调用通用性参数设置里的数据
    if CRRspecialmethod.objects.filter(cRRspecial=id2):
        lowvalue = CRRspecialmethod.objects.get(cRRspecial=id2).lowvalue  # 回收率下限
        upvalue = CRRspecialmethod.objects.get(cRRspecial=id2).upvalue  # 回收率上限

    else:
        general_1 = General.objects.get(name="通用性项目")
        crr_general = CRRgeneral.objects.get(general=general_1)
        lowvalue = CRRgeneralmethod.objects.get(cRRgeneral=crr_general).lowvalue  # 回收率下限
        upvalue = CRRgeneralmethod.objects.get(cRRgeneral=crr_general).upvalue  # 回收率上限

    # 后台管理系统-各项目参数设置-PT指标设置里找到是否设置了每个化合物的单位
    special = Special.objects.get(project=project)
    pt_special = PTspecial.objects.get(special=special)
    pt_accept = PTspecialaccept.objects.filter(pTspecial=pt_special)

    # 后台管理系统中设置的本项目化合物名称
    PTnorm = []  
    for i in pt_accept:
        PTnorm.append(i.norm)

    # 后台管理系统中设置的本项目每个化合物单位
    Unitlist = []
    for i in pt_accept:
        Unitlist.append(i.unit)

    #  第二步:开始文件读取

    '''
    注释1:csv,txt,xlsx,docx 4种格式数据读取完毕后,需要生成两个列表CRRsamplename和CRRconc,数据格式如下：
    1 稀释倍数samplename列表CRRsamplename(两个化合物内含两个列表):[ ['1times', '1times', '1times', '1times', '1times', 
    '2times', '2times', '2times', '2times', '2times', '4times', '4times', '4times', '4times', '4times', 
    '5times', '5times', '5times', '5times', '5times'], 
    ['1times', '1times', '1times', '1times', '1times', '2times', '2times', '2times', '2times', '2times', 
    '4times', '4times', '4times', '4times', '4times', '5times', '5times', '5times', '5times', '5times'] ]

    2 检测浓度列表CRRconc(两个化合物内含两个列表,与稀释倍数列表一一对应):[['120.16', '120.77', '121.83', '118.4', '120.66', '59.12', '57.49', '57.77', 
    '58.15', '58.46', '28.64', '29.39', '28.99', '28.9', '29.96', '23.07', '22.5', '23.5', '22.99', '22.93'], 
    ['115.65', '117.73', '114.9', '114.66', '117.59', '56.81', '55.89', '55.85', '55.34', '57.54', '28.15', 
    '28.61', '28.52', '27.75', '28.12', '22.67', '23.32', '22.26', '22.64', '22.82']]
    '''

    #  创新第二步需要生成的结果容器
    CRRsamplename = []  # 稀释倍数samplename列表
    CRRconc = []  # 检测浓度列表

    for file in files:
        if platform == "液质":
            if manufacturers == "Agilent":
                # 1 读取csv文件（Agilent）
                file.seek(0)
                file_data = file.read().decode('utf-8')
                lines = file_data.split('\r\n')
                for i in range(len(lines)):
                    if len(lines[i]) != 0:
                        # 以逗号分隔字符串,但忽略双引号内的逗号
                        lines[i] = re.split(r',\s*(?![^"]*\"\,)', lines[i])
                    else:
                        lines[i] = re.split(r',\s*(?![^"]*\"\,)', lines[i])
                        del lines[i]  # 最后一行如为空行，则删除该元素
                
                # 药物浓度平台的数据导出格式较不同
                if Detectionplatform=="治疗药物检测平台":
                    norm = []  # 化合物列表
                    for j in range(len(lines[0])):  # 从第一行开始
                        if " Results" in lines[0][j]:
                            # 若原始字符串中含有','，切割完后首位会多出一个'"',需去除
                            if lines[0][j].split(" Results")[0][0] != '"':
                                norm.append(lines[0][j].split(" Results")[0])
                            else:
                                norm.append(lines[0][j].split(" Results")[0][1:])

                    # 定性指标的索引
                    dxindex = []
                    for i in range(len(norm)):
                        if "2" in norm[i]:
                            dxindex.append(i)

                    # 删除norm中的定性指标和定量指标的"1"
                    norm=[norm[i] for i in range(len(norm)) if("2" not in norm[i])]
                    norm=[norm[i].split("1")[0] for i in range(len(norm)) if("1" in norm[i])]
                    
                    # 从第二行确定实验号（Sample Name）,浓度（Exp. Conc.）的索引
                    nameindex = 0  # 实验号索引
                    concindex = []  # 浓度索引列表（可能不止一个化合物，因此需要把索引放在一个列表里）
                    for j in range(len(lines[1])):  # 从第二行开始
                        if lines[1][j] == "Name":
                            nameindex = j
                        elif lines[1][j] == "Final Conc.":
                            concindex.append(j)
                    
                    # 删除concindex中的dxindex
                    concindex=[concindex[i] for i in range(len(concindex)) if(i not in dxindex)]
                    print(concindex)
                
                else:
                    # 从第一行确定化合物名称(含有"-Q Results"),并添加进入化合物列表
                    norm = []  # 化合物列表
                    for j in range(len(lines[0])):  # 从第一行开始
                        if "-Q Results" in lines[0][j]:
                            # 若原始字符串中含有','，切割完后首位会多出一个'"',需去除
                            if lines[0][j].split("-Q")[0][0] != '"':
                                norm.append(lines[0][j].split("-Q")[0])
                            else:
                                norm.append(lines[0][j].split("-Q")[0][1:])

                    # 从第二行确定实验号（Sample Name）,浓度（Exp. Conc.）的索引
                    nameindex = 0  # 实验号索引
                    concindex = []  # 浓度索引列表（可能不止一个化合物，因此需要把索引放在一个列表里）
                    for j in range(len(lines[1])):  # 从第二行开始
                        if lines[1][j] == "Name":
                            nameindex = j
                        elif lines[1][j] == "Calc. Conc.":
                            concindex.append(j)

                # 匹配原始数据中与稀释倍数相关(实验号前含有"CRR")的行
                for j in range(len(norm)):
                    norm_samplenamelist = []  # 一个化合物的稀释倍数Sample Name列表
                    norm_conclist = []  # 一个化合物的检测浓度列表
                    for i in range(len(lines)):  # 循环原始数据中的每一行
                        if "CRR" in lines[i][nameindex]:
                            norm_samplenamelist.append(lines[i][nameindex][0:10])
                            norm_conclist.append(lines[i][concindex[j]])

                    CRRsamplename.append(norm_samplenamelist)
                    CRRconc.append(norm_conclist)

            elif manufacturers == "Waters":
                # 内标标识
                ISlist = ["D3", "D4", "D5", "D6", "D7", "D8", "dx", "d4", "d8"]

                data = xlrd.open_workbook(filename=None, file_contents=file.read())  # 读取表格
                file_data = data.sheets()[0]
                nrows = file_data.nrows
                ncols = file_data.ncols

                norm = []  # 化合物列表
                Compound_row =[] # 含有“Compound”关键词的所在行(包含内标)
                norm_row = []  # 实际化合物所在行(不包含内标)
                for i in range(nrows):
                    if "Compound" in str(file_data.row_values(i)[0]) and ":" in str(file_data.row_values(i)[0]):
                        Compound_row.append(i)  

                    # 判断是否含有内标标识
                    if all(j not in str(file_data.row_values(i)[0]) for j in ISlist):
                        if "Compound" in str(file_data.row_values(i)[0]) and ":" in str(file_data.row_values(i)[0]):  # 如果某一行第一列含有关键词"Compound"，则该行中含有化合物名称，化合物名称在：后
                            norm.append(file_data.row_values(i)[0].split(":")[1].strip()) # strip()的作用是去除前后空格
                            norm_row.append(i)     

                nameindex = 0
                concindex = 0
                # 第一个化合物表格确定samplename和浓度所在列，norm_row[0]为第一个化合物所在行，+2是该化合物表格位于该化合物所在行的下两行
                for i in range(len(file_data.row_values(norm_row[0]+2))):
                    if file_data.row_values(norm_row[0]+2)[i] == "ID":
                        nameindex = i
                    elif "nmol/L" in file_data.row_values(norm_row[0]+2)[i]:
                        concindex = i

                # 未准确设置表头列名,直接返回并提示!
                if nameindex==0 or concindex==0:
                    error_message="未准确设置表头列名!"
                    return {"error_message":error_message}

                for j in range(len(norm)):
                    norm_samplenamelist = []  # 一个化合物的稀释倍数Sample Name列表
                    norm_conclist = []  # 一个化合物的检测浓度列表
                    if j < len(norm)-1:  # 如果不是最后一个化合物，索引为该化合物所在行到后一个化合物所在行
                        for i in range(norm_row[j], norm_row[j+1]):
                            # nameindex为样品名称索引，concindex为浓度索引
                            if "times" in file_data.row_values(i)[nameindex]:
                                # 第一种情况：稀释倍数为个位
                                if len(file_data.row_values(i)[nameindex].split("times")[0])==5:
                                    norm_samplenamelist.append(file_data.row_values(i)[nameindex][0:10]) 
                                # 第二种情况：稀释倍数为两位
                                if len(file_data.row_values(i)[nameindex].split("times")[0])==6:
                                    norm_samplenamelist.append(file_data.row_values(i)[nameindex][0:11]) 

                                norm_conclist.append(float(file_data.row_values(i)[concindex]))

                    else:  # 如果是最后一个化合物，索引为该化合物所在行到总行数
                        for i in range(norm_row[j], nrows):
                            # nameindex为样品名称索引，concindex为浓度索引
                            if "times" in file_data.row_values(i)[nameindex]:
                                # 第一种情况：稀释倍数为个位
                                if len(file_data.row_values(i)[nameindex].split("times")[0])==5:
                                    norm_samplenamelist.append(file_data.row_values(i)[nameindex][0:10]) 
                                # 第二种情况：稀释倍数为两位
                                if len(file_data.row_values(i)[nameindex].split("times")[0])==6:
                                    norm_samplenamelist.append(file_data.row_values(i)[nameindex][0:11]) 
                                norm_conclist.append(float(file_data.row_values(i)[concindex]))

                    CRRsamplename.append(norm_samplenamelist)
                    CRRconc.append(norm_conclist)

            elif manufacturers == "Thermo":
                data = xlrd.open_workbook(filename=None, file_contents=file.read())  # 读取表格
                norm = []  # Thermo的原始数据格式为一个化合物一个sheet,获取每个sheet的名字,即为化合物名称
                sheetindex = []  # 需要的化合物所在sheet索引列表
                for index in range(len(data.sheet_names())):
                    norm.append(data.sheet_names()[index].split("Sheet1")[1])
                    sheetindex.append(index)

                # 循环读取每个sheet工作表,即为每个化合物的表
                for index in range(len(sheetindex)):
                    file_data = data.sheets()[sheetindex[index]]
                    nrows = file_data.nrows
                    ncols = file_data.ncols

                    # 第一行确定samplename和浓度所在列
                    nameindex = 0
                    concindex = 0
                    for i in range(len(file_data.row_values(0))):
                        if file_data.row_values(8)[i] == "Raw File Name":
                            nameindex = i
                        elif file_data.row_values(8)[i] == "Calculated Amount":
                            concindex = i

                    # 匹配原始数据中与稀释倍数相关(实验号前含有"CRR")的行
                    norm_samplenamelist = []  # 一个化合物的稀释倍数Sample Name列表
                    norm_conclist = []  # 一个化合物的检测浓度列表
                    for i in range(nrows):  # 循环原始数据中的每一行
                        if "times" in file_data.row_values(i)[nameindex]:
                            # 第一种情况：稀释倍数为个位
                            if len(file_data.row_values(i)[nameindex].split("times")[0])==5:
                                norm_samplenamelist.append(file_data.row_values(i)[nameindex][0:10]) 
                            # 第二种情况：稀释倍数为两位
                            elif len(file_data.row_values(i)[nameindex].split("times")[0])==6:
                                norm_samplenamelist.append(file_data.row_values(i)[nameindex][0:11]) 
                            norm_conclist.append(file_data.row_values(i)[concindex])

                    CRRsamplename.append(norm_samplenamelist)
                    CRRconc.append(norm_conclist)


            elif manufacturers == "岛津":
                content = []
                for line in file:
                    try:
                        content.append(line.decode("UTF-8").replace("\r\n", "").split("\t"))  # windows下
                    except:
                        content.append(line.decode("GB2312").replace("\r\n", "").split("\t")) # linux下

                nameindex = 0
                concindex = 0  # 浓度索引，岛津的数据格式决定每个化合物的浓度所在列一定是同一列
                norm = []  # 化合物列表
                norm_row = []  # 化合物所在行

                for i in range(len(content[2])):  # 第二行确定samplename和浓度所在列
                    if content[2][i] == "数据文件名":
                        nameindex = i
                    elif content[2][i] == "浓度":
                        concindex = i

                # 未准确设置表头列名,直接返回并提示!
                if nameindex==0 or concindex==0:
                    error_message="未准确设置表头列名!"
                    return {"error_message":error_message}

                for i in range(len(content)):
                    if content[i][0] == "Name":  # 如果某一行第一列为"Name"，则该行第二列为化合物名称
                        # # 若化合物名称后含有“-”，需切除
                        # if "-" in content[i][1]:
                        #     norm.append(content[i][1].split("-")[0])
                        # else:
                        #     norm.append(content[i][1])
                        norm.append(content[i][1])
                        norm_row.append(i)

                # 匹配原始数据中与稀释倍数相关(实验号后含有"times")的行
                for j in range(len(norm)):
                    norm_samplenamelist = []  # 一个化合物的稀释倍数Sample Name列表
                    norm_conclist = []  # 一个化合物的检测浓度列表
                    if j < len(norm)-1:  # 如果不是最后一个化合物，索引为该化合物所在行到后一个化合物所在行
                        for i in range(norm_row[j], norm_row[j+1]):
                            # nameindex为样品名称索引，concindex为浓度索引
                            if "times" in content[i][nameindex]:
                                # 第一种情况：稀释倍数为个位
                                if len(content[i][nameindex].split("times")[0])==5:
                                    norm_samplenamelist.append(content[i][nameindex][0:10]) 
                                # 第二种情况：稀释倍数为两位
                                elif len(content[i][nameindex].split("times")[0])==6:
                                    norm_samplenamelist.append(content[i][nameindex][0:11]) 

                                norm_conclist.append(float(content[i][concindex]))

                    else:  # 如果是最后一个化合物，索引为该化合物所在行到总行数
                        for i in range(norm_row[j], len(content)):
                            # nameindex为样品名称索引，concindex为浓度索引
                            if "times" in content[i][nameindex]:
                                # 第一种情况：稀释倍数为个位
                                if len(content[i][nameindex].split("times")[0])==5:
                                    norm_samplenamelist.append(content[i][nameindex][0:10]) 
                                # 第二种情况：稀释倍数为两位
                                elif len(content[i][nameindex].split("times")[0])==6:
                                    norm_samplenamelist.append(content[i][nameindex][0:11]) 
                                norm_conclist.append(float(content[i][concindex]))

                    CRRsamplename.append(norm_samplenamelist)
                    CRRconc.append(norm_conclist)

            elif manufacturers == "AB":
                # 定义化合物列表，列表统一命名为norm
                norm = normAB

                # 获取上传的文件
                file_data = Document(file)

                # 每个表格最上方的标题内容列表，含有母离子和子离子的信息。需依此及母离子和子离子列表判断table索引
                paragraphs = []

                # 若标题长度为0或为换行等，不添加进入标题内容列表
                for p in file_data.paragraphs:
                    if len(p.text) != 0 and p.text != "\n" and len(p.text.strip()) != 0:
                        paragraphs.append(p.text.strip())

                # 确定table索引，母离子和子离子都与后台管理系统中设置的数值相同才证明是需要读取的定量表格
                tableindex = []
                if "血浆氨基酸" in project:
                    for i in range(len(paragraphs)):
                        for j in range(len(ZP_Method_precursor_ion)):
                            if ZP_Method_precursor_ion[j] in paragraphs[i] and ZP_Method_product_ion[j] in paragraphs[i]:
                                print(paragraphs[i]+" "+ZP_Method_precursor_ion[j]+" "+ZP_Method_product_ion[j])

                                # 血浆氨基酸48项第3个化合物定量表格的母离子和子离子“存在于”第18个化合物定性表格的母离子和子离子中，无需添加
                                if "18-Aad 2" not in paragraphs[i]:
                                    tableindex.append(2*i+1)
                            
                else:
                    for i in range(len(paragraphs)):
                        for j in range(len(ZP_Method_precursor_ion)):
                            if ZP_Method_precursor_ion[j] in paragraphs[i] and ZP_Method_product_ion[j] in paragraphs[i]:
                                tableindex.append(2*i+1)

                tables = file_data.tables #获取文件中的表格集
                
                # 循环定量表格的索引
                for k in range(len(tableindex)):

                    # 获取文件中的定量表格
                    tablequantify = tables[tableindex[k]] 

                    # 先把表格里的所有数据取出来放进一个列表中，读取速度会比直接读表格快很多
                    cells=tablequantify._cells
                    ROWS=len(tablequantify.rows)
                    COLUMNS=len(tablequantify.columns)
                    rowdatalist=[] #每一行的数据
                    rowdatagatherlist=[] #大列表，包含每一行的数据

                    for i in range(ROWS*COLUMNS):
                        text=cells[i].text.replace("\n","")
                        text=text.strip() #去除空白符
                        if i % COLUMNS != 0 or i == 0: 
                            rowdatalist.append(text)
                        else:
                            rowdatagatherlist.append(rowdatalist)
                            rowdatalist=[]
                            rowdatalist.append(text)
                    rowdatagatherlist.append(rowdatalist)

                    nameindex = 0
                    concindex = 0

                    # 读取表格的第一行的单元格,判断实验号和浓度索引
                    for i in range(len(rowdatagatherlist[0])):
                        if rowdatagatherlist[0][i] == "Sample Name":
                            nameindex = i
                        elif "Calculated Conc" in rowdatagatherlist[0][i]:
                            concindex = i

                    # 匹配原始数据中与稀释倍数相关(实验号前含有"CRR")的行
                    norm_samplenamelist = []  # 一个化合物的稀释倍数Sample Name列表
                    norm_conclist = []  # 一个化合物的检测浓度列表
                    for i in range(len(rowdatagatherlist)):
                        if "CRR" in rowdatagatherlist[i][nameindex]:
                            # CRR实验号命名:"CRR-1times-1"。若稀释倍数为个位,[0:10]取出"CRR-1times";若稀释倍数为两位,[0:11]取出"CRR-10times"

                            # 第一种情况：稀释倍数为个位
                            if len(rowdatagatherlist[i][nameindex].split("times")[0])==5:
                                norm_samplenamelist.append(rowdatagatherlist[i][nameindex][0:10]) 
                            # 第二种情况：稀释倍数为两位
                            if len(rowdatagatherlist[i][nameindex].split("times")[0])==6:
                                norm_samplenamelist.append(rowdatagatherlist[i][nameindex][0:11]) 
                            norm_conclist.append(rowdatagatherlist[i][concindex])

                    CRRsamplename.append(norm_samplenamelist)
                    CRRconc.append(norm_conclist)

        elif platform == "液相":
            if manufacturers == "Agilent":

                # .xlsx格式
                data = xlrd.open_workbook(filename=None, file_contents=file.read())  # 读取表格
                file_data = data.sheets()[0]
                nrows = file_data.nrows
                ncols = file_data.ncols

                norm = []  # 化合物列表
                norm_row = []  # 化合物所在行
                for j in range(nrows):
                    # 如果某一行的第一个元素为“化合物”，则添加第三个元素进入化合物列表
                    if file_data.row_values(j)[0] == "化合物:":
                        norm.append(file_data.row_values(j)[2])
                        norm_row.append(j)

                nameindex = 0
                concindex = 0

                # 第一个化合物表格确定samplename和浓度所在列，norm_row[0]为第一个化合物所在行，+2是该化合物表格位于该化合物所在行的下两行
                for i in range(len(file_data.row_values(norm_row[0]+1))):
                    if file_data.row_values(norm_row[0]+1)[i] == "样品名称":
                        nameindex = i
                    elif "含量" in file_data.row_values(norm_row[0]+1)[i]:
                        concindex = i

                # 未准确设置表头列名,直接返回并提示!
                if nameindex==0 or concindex==0:
                    error_message="未准确设置表头列名!"
                    return {"error_message":error_message}

                for j in range(len(norm)):
                    norm_samplenamelist = []  # 一个化合物的稀释倍数Sample Name列表
                    norm_conclist = []  # 一个化合物的检测浓度列表
                    if j < len(norm)-1:  # 如果不是最后一个化合物，索引为该化合物所在行到后一个化合物所在行
                        for i in range(norm_row[j], norm_row[j+1]):
                            # nameindex为样品名称索引，concindex为浓度索引
                            if "times" in file_data.row_values(i)[nameindex]:
                                # 第一种情况：稀释倍数为个位
                                if len(file_data.row_values(i)[nameindex].split("times")[0])==5:
                                    norm_samplenamelist.append(file_data.row_values(i)[nameindex][0:10]) 
                                # 第二种情况：稀释倍数为两位
                                if len(file_data.row_values(i)[nameindex].split("times")[0])==6:
                                    norm_samplenamelist.append(file_data.row_values(i)[nameindex][0:11]) 

                                norm_conclist.append(float(file_data.row_values(i)[concindex]))

                    else:  # 如果是最后一个化合物，索引为该化合物所在行到总行数
                        for i in range(norm_row[j], nrows):
                            # nameindex为样品名称索引，concindex为浓度索引
                            if "times" in file_data.row_values(i)[nameindex]:
                                # 第一种情况：稀释倍数为个位
                                if len(file_data.row_values(i)[nameindex].split("times")[0])==5:
                                    norm_samplenamelist.append(file_data.row_values(i)[nameindex][0:10]) 
                                # 第二种情况：稀释倍数为两位
                                if len(file_data.row_values(i)[nameindex].split("times")[0])==6:
                                    norm_samplenamelist.append(file_data.row_values(i)[nameindex][0:11]) 
                                norm_conclist.append(float(file_data.row_values(i)[concindex]))

                    CRRsamplename.append(norm_samplenamelist)
                    CRRconc.append(norm_conclist)

        ########文件读取完毕#######

    #  第三步:文件读取完毕后的操作

    '''
    注释2:需要生成一个字典CRR_dict,数据格式如下
    CRR_dict={'化合物1': [['原样', 120.16, 120.77, 121.83, 118.4, 120.66, '120.36', '1.0427'],
    ['2', 59.12, 57.49, 57.77, 58.15, 58.46, '58.198', '1.0886'], 
    ['4', 28.64, 29.39, 28.99, 28.9, 29.96, '29.176', '1.7632'], 
    ['5', 23.07, 22.5, 23.5, 22.99, 22.93, '22.998', '1.5524']], 

    '化合物2': [['原样', 115.65, 117.73, 114.9, 114.66, 117.59, '116.11', '1.2624'],
    ['2', 56.81, 55.89, 55.85, 55.34, 57.54, '56.286', '1.5613'], 
    ['4', 28.15, 28.61, 28.52, 27.75, 28.12, '28.230', '1.2238'], 
    ['5', 22.67, 23.32, 22.26, 22.64, 22.82, '22.742', '1.6860']]}
    '''

    print("CRRsamplename:%s" % (CRRsamplename))
    print("CRRconc:%s" % (CRRconc))

    # 如稀释倍数表格顺序不是按照顺序排列，需反转列表
    if "1times" not in CRRsamplename[0][0]:
        CRRsamplename2 = []
        for lst in CRRsamplename:
            CRRsamplename2.append(Reverse(lst))

        CRRconc2 = []
        for lst in CRRconc:
            CRRconc2.append(Reverse(lst))

        print("CRRsamplename2:%s" % (CRRsamplename2))
        print("CRRconc2:%s" % (CRRconc2))

        CRRsamplename = CRRsamplename2
        CRRconc = CRRconc2

    # 创新第三步需要生成的结果容器
    CRR_dict = {}
    CRR_judgenum = 0

    # CRRsamplename列表去重,并按稀释倍数由小到大排序
    CRRsamplename_distinct = []
    for i in CRRsamplename:
        middle_list = []
        for j in i:
            if j not in middle_list:
                middle_list.append(j)
        # middle_list.sort() # 按稀释倍数由小到大排序
        CRRsamplename_distinct.append(middle_list)
    
    print("CRRsamplename_distinct:%s" % (CRRsamplename_distinct))

    # 计算每个化合物原样的平均值，后面计算回收率需要用到
    CRR_1timesmean = []  # 每个化合物原样的平均值列表
    for j in range(len(CRRsamplename_distinct)):
        norm_conclist = []  # 每个化合物的数据列表
        for i in range(len(CRRsamplename_distinct[j])):
            if "1times" in CRRsamplename_distinct[j][i]:
                CRR_1times_conc = []  # 单独检测浓度列表，方便计算均值
                for k1 in range(5*i, 5*(i+1)):  # 检测浓度固定是5个，位于第i位的稀释倍数对应的浓度索引为5*i,5*(i+1)
                    CRR_1times_conc.append(float(effectnum(CRRconc[j][k1], digits)))  # 添加浓度
                CRR_1timesmean.append(new_round(np.mean(CRR_1times_conc), 2))  # 添加均值

    print("CRR_1timesmean:%s" % (CRR_1timesmean))

    for j in range(len(CRRsamplename_distinct)):
        norm_conclist = []  # 每个化合物的数据列表
        for i in range(len(CRRsamplename_distinct[j])):
            if "1times" in CRRsamplename_distinct[j][i]:
                CRR_1times = []  # 原样列表
                CRR_1times.append("原样")
                CRR_1times_conc = []  # 单独检测浓度列表，方便计算均值和CV
                for k1 in range(5*i, 5*(i+1)):  # 检测浓度固定是5个，位于第i位的稀释倍数对应的浓度索引为5*i,5*(i+1)
                    # 添加浓度,要放入html中的数据不要转换为数值,否则末尾0消失
                    CRR_1times.append(effectnum(CRRconc[j][k1], digits))
                    # 需要计算均值和CV的数据需要转为数值
                    CRR_1times_conc.append(float(effectnum(CRRconc[j][k1], digits)))
                CRR_1times.append(new_round(np.mean(CRR_1times_conc), 2))  # 添加均值
                CRR_1times.append(new_round(np.std(CRR_1times_conc, ddof=1)/np.mean(CRR_1times_conc)*100, 2))  # 添加CV

                for k2 in range(5*i, 5*(i+1)):  # 添加回收率
                    CRR_1times.append("/")  # 原样没有回收率,添加反斜杠

                norm_conclist.append(CRR_1times)

            else:
                CRR_othertimes = []
                # 稀释倍数实验号命名规则:CRR-1times-1

                # 第一种情况：稀释倍数为个位
                if len(CRRsamplename_distinct[j][i].split("times")[0])==5:
                    times = CRRsamplename_distinct[j][i][4:5]

                # 第二种情况：稀释倍数为两位
                if len(CRRsamplename_distinct[j][i].split("times")[0])==6:
                    times = CRRsamplename_distinct[j][i][4:6]

                CRR_othertimes.append(times)
                CRR_othertimes_conc = []  # 单独检测浓度列表，方便计算均值和CV
                for k1 in range(5*i, 5*(i+1)):
                    CRR_othertimes.append(effectnum(CRRconc[j][k1], digits))  # 添加浓度
                    CRR_othertimes_conc.append(float(effectnum(CRRconc[j][k1], digits)))
                CRR_othertimes.append(new_round(np.mean(CRR_othertimes_conc), 2))  # 添加均值
                CRR_othertimes.append(new_round(np.std(CRR_othertimes_conc, ddof=1)/np.mean(CRR_othertimes_conc)*100, 2))  # 添加CV
                for k2 in range(5*i, 5*(i+1)):
                    Recovery_rate = new_round(float(effectnum(CRRconc[j][k2], digits))*float(times)/float(CRR_1timesmean[j])*100, 2)  # 回收率
                    if float(Recovery_rate) < lowvalue or float(Recovery_rate) > upvalue:
                        CRR_judgenum += 1
                        CRR_othertimes.append(Recovery_rate)  # 添加每个稀释倍数的回收率
                    else:
                        CRR_othertimes.append(Recovery_rate)

                norm_conclist.append(CRR_othertimes)

        CRR_dict[norm[j]] = norm_conclist

    if CRR_judgenum == 0:
        insert_list = []
        for key, value in CRR_dict.items():
            for i in value:
                insert_list.append(CRR(reportinfo=reportinfo, norm=key, Dilution=i[0], test_conc1=i[1], test_conc2=i[2], test_conc3=i[3],
                                       test_conc4=i[4], test_conc5=i[5], mean_conc=i[6], cv_conc=i[7], calresults=i[8]+","+i[9]+","+i[10]+","+i[11]+","+i[12]))

        CRR.objects.bulk_create(insert_list)

    else:
        if Detectionplatform=="研发与创新平台":
            insert_list = []
            for key, value in CRR_dict.items():
                for i in value:
                    insert_list.append(CRR(reportinfo=reportinfo, norm=key, Dilution=i[0], test_conc1=i[1], test_conc2=i[2], test_conc3=i[3],
                                        test_conc4=i[4], test_conc5=i[5], mean_conc=i[6], cv_conc=i[7], calresults=i[8]+","+i[9]+","+i[10]+","+i[11]+","+i[12]))

            CRR.objects.bulk_create(insert_list)

    return {"CRR_dict": CRR_dict, "Unit": Unit,"lowvalue":lowvalue,"upvalue":upvalue}

# CRR数据关联进入最终报告
def related_CRR(id,unit):

    # 第一步：后台描述性内容数据提取
    # 1 根据id找到项目
    project = ReportInfo.objects.get(id=id).project

    # 2 优先查找特殊参数设置里是否有数据，如有就调用，没有则调用通用性参数设置里的数据
    #特殊参数设置描述性内容
    textlist_special = []
    try:
        special_1 = Special.objects.get(project=project) 
        special_2 = CRRspecial.objects.get(special=special_1)           
        if CRRspecialtexts.objects.filter(cRRspecial=special_2).count()>0:
            text_special = CRRspecialtexts.objects.filter(cRRspecial=special_2)  
            for i in text_special:
                textlist_special.append(i.text)
    except:
        pass

    # 3 通用数据抓取
    # 描述性内容
    textlist_general = [] 
    general_1 = General.objects.get(name="通用性项目") #通用参数设置描述性内容
    general_2 = CRRgeneral.objects.get(general=general_1)
    text_general = CRRgeneraltexts.objects.filter(cRRgeneral=general_2)      
    for i in text_general:
        textlist_general.append(i.text)



    # 查找是否单独设置了每个化合物的有效位数
    DIGITS_TABLE = Special.objects.get(project=project)
    pt_special = PTspecial.objects.get(special=DIGITS_TABLE)
    pt_accept = PTspecialaccept.objects.filter(pTspecial=pt_special)
    Digitslist = []  # 每个化合物有效位数列表
    Digitsdict = {}  # 每个化合物有效位数字典

    for i in pt_accept:
        Digitslist.append(i.digits)

    if Digitslist == [] or Digitslist[0] == "":  # 如果全部没设置或者只是单位没设置
        pass
    else:
        for i in pt_accept:
            Digitsdict[i.norm] = i.digits

    # 第二步：报告数据提取

    '''
    注释:需要生成一个字典CRR_endreport_dict,数据格式如下
    CRR_endreport_dict={'化合物1': [['原样', 120.16, 120.77, 121.83, 118.4, 120.66, '120.36', '1.0427'],
    ['2', 59.12, 57.49, 57.77, 58.15, 58.46, '58.198', '1.0886'], 
    ['4', 28.64, 29.39, 28.99, 28.9, 29.96, '29.176', '1.7632'], 
    ['5', 23.07, 22.5, 23.5, 22.99, 22.93, '22.998', '1.5524']], 

    '化合物2': [['原样', 115.65, 117.73, 114.9, 114.66, 117.59, '116.11', '1.2624'],
    ['2', 56.81, 55.89, 55.85, 55.34, 57.54, '56.286', '1.5613'], 
    ['4', 28.15, 28.61, 28.52, 27.75, 28.12, '28.230', '1.2238'], 
    ['5', 22.67, 23.32, 22.26, 22.64, 22.82, '22.742', '1.6860']]}
    '''

    # 定义需要生成的字典
    CRR_dict = {}  # 最终需要的字典

    try:
        # 1 基础数据抓取
        CRR_data = CRR.objects.filter(reportinfo_id=id)

        CRR_norm = []  # 待测物质列表
        for i in CRR_data:
            if i.norm not in CRR_norm:
                CRR_norm.append(i.norm)

        CRR_range = []  # 稀释倍数列表，第三步要用，方便找到最大稀释倍数
        for i in CRR_norm:
            middle_list = []  # 每个化合物的数据列表
            Dilutionlist = []  # 每个化合物的稀释倍数列表，需要加到CRR_range中
            middle_table = CRR.objects.filter(reportinfo_id=id, norm=i)  # 每个待测物质的数据表          
            for item in middle_table:
                # 没有为每个化合物单独设置有效位数，则调用通用性设置
                if Digitsdict == {} or list(Digitsdict.values())[0] == None:
                    if item.Dilution != "原样":
                        Dilutionlist.append(float(item.Dilution))

                    rowlist = []  # 每一行的小列表
                    rowlist.append(item.Dilution)
                    rowlist.append(item.test_conc1)
                    rowlist.append(item.test_conc2)
                    rowlist.append(item.test_conc3)
                    rowlist.append(item.test_conc4)
                    rowlist.append(item.test_conc5)

                    rowlist.append(item.calresults.split(',')[0])
                    rowlist.append(item.calresults.split(',')[1])
                    rowlist.append(item.calresults.split(',')[2])
                    rowlist.append(item.calresults.split(',')[3])
                    rowlist.append(item.calresults.split(',')[4])

                    rowlist.append(item.mean_conc)
                    rowlist.append(item.cv_conc)
                    middle_list.append(rowlist)

                # 为每个化合物单独设置了有效位数，则调用每个化合物的设置
                else:
                    if item.Dilution != "原样":
                        Dilutionlist.append(float(item.Dilution))

                    rowlist = []
                    rowlist.append(item.Dilution)
                    rowlist.append(effectnum(item.test_conc1, Digitsdict[i]))
                    rowlist.append(effectnum(item.test_conc2, Digitsdict[i]))
                    rowlist.append(effectnum(item.test_conc3, Digitsdict[i]))
                    rowlist.append(effectnum(item.test_conc4, Digitsdict[i]))
                    rowlist.append(effectnum(item.test_conc5, Digitsdict[i]))

                    rowlist.append(item.calresults.split(',')[0])
                    rowlist.append(item.calresults.split(',')[1])
                    rowlist.append(item.calresults.split(',')[2])
                    rowlist.append(item.calresults.split(',')[3])
                    rowlist.append(item.calresults.split(',')[4])

                    rowlist.append(item.mean_conc)
                    rowlist.append(item.cv_conc)
                    middle_list.append(rowlist)

            CRR_dict[i] = middle_list
            CRR_range.append(Dilutionlist)

        # 第三步：临床可报告范围数据提取
        # 找到对应化合物AMR的上下限
        AMR_theoryconc = []  # AMR理论浓度列表，方便提取每个化合物AMR的上下限
        for i in CRR_norm:
            data_AMR = AMR.objects.filter(reportinfo_id=id, norm=i)  # AMR每个待测物质的数据表
            if data_AMR.exists():
                Dilutionlist = []  # 每个化合物AMR理论浓度列表
                for item in data_AMR:
                    Dilutionlist.append(float(item.therory_conc))
                AMR_theoryconc.append(Dilutionlist)
            else:
                pass

        Dilution = []
        print("CRR_dict")
        print(CRR_dict)
        print(CRR_range)
        print(AMR_theoryconc)
        print(CRR_norm)

        # 需进行判断，如果CRR_range或AMR_theoryconc为空列表，其后所有代码均不执行
        if len(CRR_range)!=0 and len(AMR_theoryconc)!=0:
            for i in CRR_range[0]:
                Dilution.append(str(int(i)))
            
            CRR_conclusion1 = "、".join(Dilution)

            CRR_conclusion2 = "按最大稀释倍数" + str(int(max(CRR_range[0])))+"倍计算，临床可报告范围为："

            # 多指标项目，临床可报告范围结论需用表格展示
            if len(CRR_norm)>1:
                
                CRR_conclusion_dict={}

                for i in range(len(CRR_norm)):
                    middle_list=[]
                    middle_list.append(str(min(AMR_theoryconc[i])))
                    middle_list.append(str(int(max(CRR_range[0]))*max(AMR_theoryconc[i])))
                    CRR_conclusion_dict[CRR_norm[i]]=middle_list


                    # # 最后一个化合物后无需加逗号
                    # if i!=len(CRR_norm)-1:
                    #     CRR_conclusion2 = CRR_conclusion2 + str(min(AMR_theoryconc[i]))+'~'+str(int(max(CRR_range[0]))*max(AMR_theoryconc[i])) + unit+"，"
                    # else:
                    #     CRR_conclusion2 = CRR_conclusion2 + str(min(AMR_theoryconc[i]))+'~'+str(int(max(CRR_range[0]))*max(AMR_theoryconc[i])) + unit

                print(CRR_conclusion_dict)

                # return最终需要的结果
                if len(textlist_special) != 0:
                    return {"CRR_dict": CRR_dict, "textlist": textlist_special, "serial": len(textlist_special)+1,
                            "CRR_conclusion1": CRR_conclusion1, "CRR_conclusion2": CRR_conclusion2, "CRR_conclusion_dict": CRR_conclusion_dict}

                else:
                    return {"CRR_dict": CRR_dict, "textlist": textlist_general, "serial": len(textlist_general)+1,
                            "CRR_conclusion1": CRR_conclusion1, "CRR_conclusion2": CRR_conclusion2, "CRR_conclusion_dict": CRR_conclusion_dict}

            # 单一指标项目，临床可报告范围结论用文字展示
            else:
                CRR_conclusion_dict={}

                for i in range(len(CRR_norm)):
                    middle_list=[]
                    middle_list.append(str(min(AMR_theoryconc[i])))
                    middle_list.append(str(int(max(CRR_range[0]))*max(AMR_theoryconc[i])))
                    CRR_conclusion_dict[CRR_norm[i]]=middle_list

                for i in range(len(CRR_norm)):
                    CRR_conclusion2 = CRR_conclusion2 + str(min(AMR_theoryconc[i]))+'~'+str(int(max(CRR_range[0]))*max(AMR_theoryconc[i])) + unit

                # return最终需要的结果
                if len(textlist_special) != 0:
                    return {"CRR_dict": CRR_dict, "textlist": textlist_special, "serial": len(textlist_special)+1,
                            "CRR_conclusion1": CRR_conclusion1, "CRR_conclusion2": CRR_conclusion2, "CRR_conclusion_dict": CRR_conclusion_dict}

                else:
                    return {"CRR_dict": CRR_dict, "textlist": textlist_general, "serial": len(textlist_general)+1,
                            "CRR_conclusion1": CRR_conclusion1, "CRR_conclusion2": CRR_conclusion2, "CRR_conclusion_dict": CRR_conclusion_dict}
        
        else:
            CRR_conclusion1 = ""
            CRR_conclusion2 = "请先完成AMR验证后再来看稀释倍数的最终结论"
            if len(textlist_special) != 0:
                return {"CRR_dict": CRR_dict, "textlist": textlist_special, "serial": len(textlist_special)+1,
                        "CRR_conclusion1": CRR_conclusion1, "CRR_conclusion2": CRR_conclusion2}

            else:
                return {"CRR_dict": CRR_dict, "textlist": textlist_general, "serial": len(textlist_general)+1,
                        "CRR_conclusion1": CRR_conclusion1, "CRR_conclusion2": CRR_conclusion2}

        

    except: 
        pass
